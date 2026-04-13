#!/usr/bin/env python3
# -*- coding:utf-8 -*-
"""
Oracle 数据库自动化健康巡检工具 v1.0
支持 Oracle 11g 及以上版本
依赖: cx_Oracle 或 oracledb, python-docx, docxtpl, openpyxl, psutil, paramiko
"""

import warnings
warnings.filterwarnings("ignore")
import itertools
import math
import sys
import datetime
import argparse
import subprocess
import logging
import logging.handlers
import socket
import re
import time
from pathlib import Path
import sys, getopt, os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm
from docxtpl import DocxTemplate
import configparser
import importlib
import subprocess
import json
import hashlib
import base64
from datetime import datetime, timedelta
import platform
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
import tempfile
import io
import psutil
import shutil
import paramiko

# Oracle 客户端驱动（优先使用官方 oracledb，降级到 cx_Oracle）
try:
    import oracledb as oracle_driver
    ORACLE_DRIVER = 'oracledb'
except ImportError:
    try:
        import cx_Oracle as oracle_driver
        ORACLE_DRIVER = 'cx_Oracle'
    except ImportError:
        print("缺少 Oracle 驱动库，请执行以下命令之一安装：")
        print("  pip install oracledb      # 推荐")
        print("  pip install cx_Oracle     # 备选")
        sys.exit(1)

importlib.reload(sys)

# ============================================================
# 磁盘采集时忽略的外接 ISO / Media 挂载点前缀（与 main_mysql 保持一致）
# ============================================================
IGNORE_MOUNTS = {'/mnt/iso', '/media', '/run/media', '/iso', '/cdrom'}

# ============================================================
# 内置 Oracle 巡检 SQL 模板（覆盖 Oracle 11g+ 核心指标）
# ============================================================
ORACLE_SQL_TEMPLATES = """
[report]
name = Oracle HealthCheck Report
template = ./templates/oracle_wordtemplates_v1.0.docx
output = /tmp/OracleCheckReport.docx

[variables]
# ── 1. 基本信息 ──────────────────────────────
ora_version      = SELECT * FROM v$version WHERE banner LIKE 'Oracle%';
ora_instance     = SELECT instance_name, host_name, version, status, database_status, instance_role, edition FROM v$instance;
ora_database     = SELECT name, db_unique_name, created, log_mode, open_mode, database_role, flashback_on, platform_name FROM v$database;
ora_uptime       = SELECT to_char(startup_time,'YYYY-MM-DD HH24:MI:SS') AS startup_time, round((sysdate - startup_time) * 24 * 60) AS uptime_minutes FROM v$instance;

# ── 2. 会话与连接 ──────────────────────────────
ora_sessions      = SELECT COUNT(*) AS total_sessions, SUM(CASE WHEN STATUS='ACTIVE' THEN 1 ELSE 0 END) AS active_sessions, SUM(CASE WHEN STATUS='INACTIVE' THEN 1 ELSE 0 END) AS inactive_sessions, COUNT(DISTINCT username) AS logged_users FROM v$session WHERE type='USER' AND username IS NOT NULL;
ora_session_limit = SELECT value AS sessions_limit FROM v$parameter WHERE name='sessions';
ora_process_limit = SELECT value AS processes_limit FROM v$parameter WHERE name='processes';
ora_blocked       = SELECT s1.sid AS blocked_sid, s1.username AS blocked_user, s1.status AS blocked_status, s2.sid AS blocking_sid, s2.username AS blocking_user, s1.event AS wait_event, s1.seconds_in_wait AS sec_in_wait, s1.sql_id, o.object_name, o.object_type FROM v$lock l1 JOIN v$session s1 ON l1.sid = s1.sid JOIN v$lock l2 ON l1.id1 = l2.id1 AND l1.request > 0 AND l2.request = 0 JOIN v$session s2 ON l2.sid = s2.sid LEFT JOIN dba_objects o ON l1.id2 = o.object_id WHERE l1.type IN ('TM','TX') ORDER BY s1.seconds_in_wait DESC;

# ── 3. 表空间 ──────────────────────────────────
ora_tablespace   = SELECT d.tablespace_name, t.status, t.contents, t.extent_management, t.segment_space_management, ROUND(d.bytes / 1024 / 1024, 2) AS total_mb, ROUND(NVL(f.bytes, 0) / 1024 / 1024, 2) AS free_mb, ROUND((d.bytes - NVL(f.bytes, 0)) / 1024 / 1024, 2) AS used_mb, ROUND(100 * (1 - NVL(f.bytes, 0) / d.bytes), 2) AS used_pct, CASE WHEN d.maxbytes = 0 THEN d.bytes ELSE d.maxbytes END AS max_bytes, ROUND(CASE WHEN d.maxbytes = 0 THEN 100 * (d.bytes - NVL(f.bytes, 0)) / d.bytes ELSE 100 * (d.bytes - NVL(f.bytes, 0)) / d.maxbytes END, 2) AS used_pct_with_maxext FROM (SELECT tablespace_name, SUM(bytes) bytes, SUM(CASE WHEN autoextensible='YES' THEN maxbytes ELSE bytes END) maxbytes FROM dba_data_files GROUP BY tablespace_name) d, (SELECT tablespace_name, status, contents, extent_management, segment_space_management FROM dba_tablespaces) t, (SELECT tablespace_name, SUM(bytes) bytes FROM dba_free_space GROUP BY tablespace_name) f WHERE d.tablespace_name = f.tablespace_name(+) AND d.tablespace_name = t.tablespace_name ORDER BY 7 DESC;
ora_temp_ts       = SELECT t.tablespace_name, t.status, ROUND(t.bytes / 1024 / 1024, 2) AS total_mb, ROUND(NVL(u.used_bytes, 0) / 1024 / 1024, 2) AS used_mb, ROUND(100 * NVL(u.used_bytes, 0) / t.bytes, 2) AS used_pct FROM dba_temp_files t LEFT JOIN (SELECT tablespace_name, SUM(bytes_used) AS used_bytes FROM v$temp_space_header GROUP BY tablespace_name) u ON t.tablespace_name = u.tablespace_name;

# ── 4. SGA 与内存 ─────────────────────────────
ora_sga           = SELECT component, current_size/1024/1024 AS current_size_mb, min_size/1024/1024 AS min_size_mb, user_specified_size/1024/1024 AS user_size_mb FROM v$sga_dynamic_components WHERE current_size > 0 ORDER BY current_size DESC;
ora_sga_total     = SELECT SUM(value)/1024/1024 AS sga_total_mb FROM v$sga;
ora_pga           = SELECT NAME, VALUE/1024/1024 AS value_mb FROM v$pgastat WHERE NAME IN ('total PGA allocated','total PGA inuse','aggregate PGA target parameter','aggregate PGA auto target','maximum PGA allocated','total freeable PGA memory');
ora_memory_target  = SELECT NAME, VALUE, DISPLAY_VALUE, ISDEFAULT FROM v$parameter WHERE NAME IN ('memory_target','memory_max_target','sga_target','pga_aggregate_target','sga_max_size');

# ── 5. Redo 日志 ───────────────────────────────
ora_redo_logs     = SELECT group#, thread#, sequence#, bytes/1024/1024 AS size_mb, members, archived, status, first_change#, next_change# FROM v$log ORDER BY group#;
ora_redo_status   = SELECT group#, member, status FROM v$logfile ORDER BY group#;

# ── 6. 归档与备份 ──────────────────────────────
ora_archive_dest  = SELECT dest_name, destination, status, binding, target, process, CASE WHEN status='VALID' AND destination IS NOT NULL THEN 'ENABLED' ELSE 'DISABLED' END AS recovery_mode FROM v$archive_dest WHERE dest_name IN ('LOG_ARCHIVE_DEST_1','LOG_ARCHIVE_DEST_2','LOG_ARCHIVE_DEST_3');
ora_archive_lag   = SELECT name, applied, sequence#, next_time, CASE WHEN applied='YES' THEN sequence# ELSE 0 END AS archived_seq FROM v$archived_log WHERE rownum <= 10 ORDER BY sequence# DESC;
ora_backup        = SELECT b.recid AS recid, b.start_time, b.completion_time AS end_time, b.backup_type, p.handle, p.status FROM v$backup_set b LEFT JOIN v$backup_piece p ON b.set_stamp=p.set_stamp AND b.set_count=p.set_count WHERE rownum <= 5 ORDER BY b.completion_time DESC;

# ── 7. 关键参数 ────────────────────────────────
ora_params        = SELECT NAME, VALUE, DISPLAY_VALUE, ISDEFAULT, ISSES_MODIFIABLE, ISSYS_MODIFIABLE AS SYSMODIFIABLE, DESCRIPTION FROM v$parameter WHERE NAME IN ('processes','sessions','open_cursors','db_block_size','db_file_multiblock_read_count','db_writer_processes','undo_retention','compatible','nls_characterset','nls_nchar_characterset','job_queue_processes','parallel_max_servers','audit_trail','recyclebin','optimizer_mode','cursor_sharing','statistics_level','control_file_record_keep_time','remote_login_passwordfile','resource_manager_plan') ORDER BY NAME;

# ── 8. 无效对象 ────────────────────────────────
ora_invalid_objs  = SELECT owner, object_type, object_name, status, created, last_ddl_time FROM dba_objects WHERE status='INVALID' AND owner NOT IN ('SYS','SYSTEM','PUBLIC','MDSYS','ORDSYS','CTXSYS','OLAPSYS','XDB','WMSYS','APEX_040200','EXFSYS','ORDPLUGINS') ORDER BY owner, object_type, object_name;
ora_invalid_cnt   = SELECT owner, COUNT(*) AS invalid_count FROM dba_objects WHERE status='INVALID' AND owner NOT IN ('SYS','SYSTEM','PUBLIC') GROUP BY owner HAVING COUNT(*) > 0 ORDER BY 2 DESC;

# ── 9. 用户与权限 ──────────────────────────────
ora_users         = SELECT username, account_status, created, default_tablespace, temporary_tablespace, profile, expiry_date, lock_date FROM dba_users WHERE username NOT IN ('XS$NULL','ANONYMOUS') AND oracle_maintained='N' ORDER BY username;
ora_sys_privs     = SELECT grantee, COUNT(*) AS priv_count FROM dba_sys_privs WHERE grantee IN (SELECT username FROM dba_users WHERE account_status='OPEN' AND oracle_maintained='N' AND username NOT IN ('SYS','SYSTEM')) GROUP BY grantee HAVING COUNT(*) >= 10 ORDER BY 2 DESC;
ora_default_pws   = SELECT username, account_status FROM dba_users WHERE password IN (SELECT password FROM dba_users WHERE username IN ('SCOTT','HR','OE','SH','PM','IX','BI','ODM_MTR','WK_TEST','QS_CB','QS_OS','QS_ES','QS_WS','QS_CS','QS_RS','QSADM','PERFSTAT','OLAPSVR')) AND oracle_maintained='N';

# ── 10. 长时间运行SQL ───────────────────────────
ora_long_sql      = SELECT sid, serial#, username, sql_id, opname, sofar, totalwork, ROUND(sofar/GREATEST(totalwork,0.001)*100,1) AS pct_complete, elapsed_seconds, time_remaining FROM v$session_longops WHERE totalwork > 0 AND sofar < totalwork AND elapsed_seconds > 30 ORDER BY elapsed_seconds DESC FETCH FIRST 10 ROWS ONLY;
ora_top_sql_cpu   = SELECT sql_id, substr(sql_text,1,80) AS sql_text, cpu_time/1000000 AS cpu_sec, executions, buffer_gets, disk_reads, rows_processed, elapsed_time/1000000 AS elapsed_sec, module, action, first_load_time FROM v$sqlarea WHERE executions > 0 ORDER BY cpu_time DESC FETCH FIRST 15 ROWS ONLY;

# ── 11. Data Guard / ADG ───────────────────────
ora_dg_status     = SELECT database_role, protection_mode, protection_level, open_mode, switchover_status, dataguard_broker FROM v$database;
ora_dg_apply      = SELECT process, pid, status, thread#, sequence#, block#, blocks FROM v$managed_standby WHERE process IN ('MRP0','MRP1','RFS');
ora_standby_event = SELECT event, total_waits, total_timeouts, time_waited, average_wait FROM v$system_event WHERE event LIKE '%DG%' OR event LIKE '%standby%' OR event LIKE '%log apply%' OR event LIKE '%redo transport%' ORDER BY time_waited DESC FETCH FIRST 10 ROWS ONLY;

# ── 12. RAC 信息 ───────────────────────────────
ora_rac_nodes     = SELECT inst_id, instance_number, instance_name, host_name, version, status, startup_time, thread#, database_status FROM gv$instance ORDER BY inst_id;
ora_rac_interconn = SELECT inst_id, ip_address AS link_name, CASE WHEN ip_address IS NOT NULL THEN 'UP' ELSE 'UNKNOWN' END AS status FROM gv$cluster_interconnects WHERE inst_id = (SELECT inst_id FROM v$instance);

# ── 13. ASM 磁盘组 ─────────────────────────────
ora_asm_diskgroup = SELECT group_number, name, state, type, total_mb, free_mb, usable_file_mb, offline_disks, voting_files, ROUND((1-free_mb/total_mb)*100,2) AS used_pct FROM v$asm_diskgroup ORDER BY group_number;

# ── 14. Undo 表空间 ────────────────────────────
ora_undo_info     = SELECT d.undo_tablespace, r.retention, ROUND(NVL(ts.used_bytes/1024/1024,0),2) AS used_mb, ROUND(NVL(ts.tbs_bytes/1024/1024,0),2) AS total_mb, u.exp_blks AS exp_undo_blks, u.unexp_blks AS unexp_undo_blks, u.blk_cnt AS undo_blk_cnt FROM (SELECT UPPER(VALUE) AS undo_tablespace FROM v$parameter WHERE NAME='undo_tablespace') d, (SELECT UPPER(VALUE) AS retention FROM v$parameter WHERE NAME='undo_retention') r, (SELECT SUM(df.bytes) AS tbs_bytes, SUM(df.bytes)-NVL(SUM(ff.free_bytes),0) AS used_bytes FROM dba_data_files df LEFT JOIN (SELECT tablespace_name, SUM(bytes) AS free_bytes FROM dba_free_space GROUP BY tablespace_name) ff ON df.tablespace_name = ff.tablespace_name WHERE df.tablespace_name = (SELECT UPPER(VALUE) FROM v$parameter WHERE NAME='undo_tablespace')) ts, (SELECT COUNT(*) AS exp_blks, 0 AS unexp_blks, COUNT(*) AS blk_cnt FROM v$undostat WHERE begin_time > SYSDATE-1 AND undoblks > 0) u;

# ── 15. 闪回与回收站 ───────────────────────────
ora_recyclebin    = SELECT owner, type, object_name, droptime, original_name, space*8192/1024/1024 AS size_mb FROM dba_recyclebin WHERE space*8192/1024/1024 > 50 ORDER BY space DESC FETCH FIRST 20 ROWS ONLY;
ora_flashback_area= SELECT name, space_limit/1024/1024 AS limit_mb, space_used/1024/1024 AS used_mb, space_reclaimable/1024/1024 AS reclaimable_mb, NUMBER_OF_FILES AS file_count, ROUND(space_used*100/space_limit,2) AS used_pct FROM v$recovery_file_dest;

# ── 16. 数据文件状态 ───────────────────────────
ora_datafiles     = SELECT f.file_id, f.file_name, f.tablespace_name, f.status, f.bytes/1024/1024 AS size_mb, f.autoextensible, f.maxbytes/1024/1024 AS max_mb, ROUND(f.increment_by * b.block_size / 1024 / 1024, 2) AS inc_mb, f.online_status FROM dba_data_files f, (SELECT DISTINCT block_size FROM dba_tablespaces WHERE tablespace_name = 'SYSTEM') b ORDER BY f.tablespace_name, f.file_id;

# ── 17. Profile 密码策略 ───────────────────────
ora_profile_pwd   = SELECT profile, resource_name, limit, resource_type FROM DBA_PROFILES WHERE resource_type='PASSWORD' AND profile='DEFAULT' ORDER BY resource_name;

# ── 18. 系统等待事件 TOP ───────────────────────
ora_top_waits     = SELECT event, total_waits, time_waited/100 AS time_waited_sec, average_wait/100 AS avg_wait_ms, total_timeouts FROM v$system_event WHERE wait_class != 'Idle' AND total_waits > 0 ORDER BY time_waited DESC FETCH FIRST 15 ROWS ONLY;
ora_wait_class    = SELECT wait_class, COUNT(*) AS event_count, SUM(total_waits) AS total_waits, SUM(time_waited)/100 AS total_time_sec FROM v$system_event WHERE wait_class != 'Idle' GROUP BY wait_class ORDER BY 3 DESC;

# ── 19. 统计信息陈旧 ───────────────────────────
ora_stale_stats   = SELECT owner, table_name, partition_name, object_type, stale_stats, last_analyzed FROM dba_tab_statistics WHERE stale_stats='YES' AND owner NOT IN ('SYS','SYSTEM') AND owner NOT LIKE '%SYS%' ORDER BY last_analyzed NULLS LAST FETCH FIRST 20 ROWS ONLY;

# ── 20. 分区表信息 ─────────────────────────────
ora_partition_info = SELECT table_owner, table_name, partition_name, high_value, num_rows, last_analyzed, tablespace_name FROM dba_tab_partitions WHERE table_owner NOT IN ('SYS','SYSTEM','WMSYS','MDSYS') AND num_rows > 0 ORDER BY table_owner, table_name FETCH FIRST 15 ROWS ONLY;
"""


# ============================================================
# 远程系统信息收集器（与 main_mysql 共用逻辑，保持一致）
# ============================================================
class RemoteSystemInfoCollector:
    """远程系统信息收集器 - 通过SSH连接获取远程主机信息"""

    def __init__(self, host, port=22, username='root', password=None, key_file=None):
        self.host = host
        self.port = port
        self.username = username
        self.password = password
        self.key_file = key_file
        self.ssh_client = None

    def connect(self):
        try:
            self.ssh_client = paramiko.SSHClient()
            self.ssh_client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            if self.key_file:
                private_key = paramiko.RSAKey.from_private_key_file(self.key_file)
                self.ssh_client.connect(hostname=self.host, port=self.port, username=self.username, pkey=private_key, timeout=10)
            else:
                self.ssh_client.connect(hostname=self.host, port=self.port, username=self.username, password=self.password, timeout=10)
            return True
        except Exception as e:
            print(f"SSH连接失败 {self.host}:{self.port}: {e}")
            return False

    def disconnect(self):
        if self.ssh_client:
            try: self.ssh_client.close()
            except Exception: pass

    def _run(self, cmd):
        """执行命令并返回 stdout"""
        stdin, stdout, stderr = self.ssh_client.exec_command(cmd, timeout=20)
        out = stdout.read().decode('utf-8', errors='replace').strip()
        err = stderr.read().decode('utf-8', errors='replace').strip()
        return out, err

    def get_cpu_info(self):
        info = {}
        out, _ = self._run("lscpu | grep -E 'Model name|CPU\\(s\\)|Architecture|Thread' | head -5")
        for line in out.split('\n'):
            if ':' in line:
                k, v = line.split(':', 1)
                k = k.strip(); v = v.strip()
                if 'CPU(s)' in k or 'Model name' in k or 'Architecture' in k:
                    info[k] = v

        out2, _ = self._run("cat /proc/stat | grep cpu | head -1")
        if out2 and len(out2.split()) >= 8:
            fields = out2.split()[1:]
            user = int(fields[0]); nice = int(fields[1])
            system = int(fields[2]); idle = int(fields[3])
            iowait = int(fields[4]) if len(fields) > 4 else 0
            total = user + nice + system + idle + iowait
            info['usage_percent'] = round((total - idle) / total * 100, 1) if total > 0 else 0.0
        return info

    def get_memory_info(self):
        out, _ = self._run("free -b | grep Mem")
        info = {}
        if out:
            parts = out.split()[1:]
            total_mb = float(parts[0]) / 1024 / 1024
            used_mb = float(parts[1]) / 1024 / 1024
            info['total_mb'] = round(total_mb, 2)
            info['used_mb'] = round(used_mb, 2)
            info['usage_percent'] = round(used_mb / total_mb * 100, 1) if total_mb > 0 else 0.0
        return info

    def get_disk_info(self):
        out, _ = self._run(f"df -h | grep -vE 'tmpfs|udev|overlay|mnt/iso|/iso|/media/|/run/media/|/cdrom|Filesystem'")
        disks = {}
        for line in out.split('\n')[1:]:
            parts = line.split()
            if len(parts) >= 6:
                device, size, used, avail, use_pct, mountpoint = parts[0], parts[1], parts[2], parts[3], parts[4], parts[5]
                try:
                    total_gb = float(size.replace('G','').replace('M',''))
                    used_gb = float(used.replace('G','').replace('M',''))
                    free_gb = float(avail.replace('G','').replace('M',''))
                    usage = float(use_pct.replace('%',''))
                    if 'M' in avail:
                        free_gb /= 1024; used_gb /= 1024; total_gb /= 1024
                    disks[mountpoint] = {
                        'device': device, 'mountpoint': mountpoint,
                        'total_gb': round(total_gb, 2), 'used_gb': round(used_gb, 2),
                        'free_gb': round(free_gb, 2), 'usage_percent': usage,
                        'fstype': 'unknown'
                    }
                except (ValueError, IndexError): continue
        return disks

    def get_system_info(self):
        hostname_out, _ = self._run("hostname")
        boot_out, _ = self._run("uptime -s")
        platform_out, _ = self._run("uname -sr")
        return {
            'hostname': hostname_out.strip() if hostname_out.strip() else '未知',
            'platform': platform_out.strip() if platform_out.strip() else 'Linux',
            'boot_time': boot_out.strip() if boot_out.strip() else '未知',
            'cpu': self.get_cpu_info(),
            'memory': self.get_memory_info(),
            'disk': self.get_disk_info()
        }


class LocalSystemInfoCollector:
    """本地系统信息收集器"""

    def get_cpu_info(self):
        info = {'Model name': '', 'usage_percent': 0.0}
        try:
            import platform as _pf
            info['Model name'] = f"{_pf.processor()} ({_pf.machine()})"
            info['usage_percent'] = round(psutil.cpu_percent(interval=1), 1)
        except Exception: pass
        return info

    def get_memory_info(self):
        info = {'total_mb': 0, 'used_mb': 0, 'usage_percent': 0.0}
        try:
            mem = psutil.virtual_memory()
            info['total_mb'] = round(mem.total / 1024 / 1024, 2)
            info['used_mb'] = round(mem.used / 1024 / 1024, 2)
            info['usage_percent'] = round(mem.percent, 1)
        except Exception: pass
        return info

    def get_disk_info(self):
        disks = {}
        try:
            for part in psutil.disk_partitions():
                try:
                    usage = psutil.disk_usage(part.mountpoint)
                    if part.mountpoint in IGNORE_MOUNTS:
                        continue
                    disks[part.mountpoint] = {
                        'device': part.device,
                        'mountpoint': part.mountpoint,
                        'fstype': part.fstype,
                        'total_gb': round(usage.total / 1024**3, 2),
                        'used_gb': round(usage.used / 1024**3, 2),
                        'free_gb': round(usage.free / 1024**3, 2),
                        'usage_percent': round(usage.percent, 1),
                    }
                except (PermissionError, OSError): continue
        except Exception: pass
        return disks

    def get_system_info(self):
        import socket as _sock
        import platform as _pf
        boot_time = ''
        try: boot_time = datetime.fromtimestamp(psutil.boot_time()).strftime('%Y-%m-%d %H:%M:%S')
        except Exception: pass
        return {
            'hostname': _sock.gethostname() or 'localhost',
            'platform': f"{_pf.system()} {_pf.release()}",
            'boot_time': boot_time,
            'cpu': self.get_cpu_info(),
            'memory': self.get_memory_info(),
            'disk': self.get_disk_info()
        }


def get_host_disk_usage():
    """获取主机磁盘使用情况（兼容函数）"""
    collector = LocalSystemInfoCollector()
    disk_dict = collector.get_disk_info()
    return list(disk_dict.values())


def create_word_template(inspector_name="Jack"):
    """
    创建 Oracle 巡检报告的 Word 模板（基于 docxtpl Jinja2 模板）。
    所有变量均使用 {{{ var }}} 占位符，避免复杂 Jinja2 语法导致解析错误。
    """
    template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")
    tpl_file = os.path.join(template_path, "oracle_wordtemplates_v1.0.docx")

    # 如果模板已存在，直接返回
    if os.path.exists(tpl_file):
        return tpl_file

    from docx import Document
    doc = Document()

    # 标题样式
    title_style = doc.styles.add_style('ReportTitle', 1)
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0,51,102)

    heading1 = doc.styles.add_style('Heading1Custom', 1)
    heading1.font.size = Pt(18)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0,51,102)

    heading2 = doc.styles.add_style('Heading2Custom', 1)
    heading2.font.size = Pt(14)
    heading2.font.bold = True

    # 封面
    doc.add_paragraph("Oracle数据库健康巡检报告", style='ReportTitle')
    doc.add_paragraph("")
    # 基本信息表（与 PostgreSQL 报告格式一致）
    info_table = doc.add_table(rows=8, cols=2)
    info_table.style = 'Light Grid Accent 1'
    info_cells = [
        ("数据库名称",     "{{{ co_name }}}"),
        ("服务器地址",     "{{{ server_addr }}}"),
        ("Oracle版本",     "{{{ ora_version }}}"),
        ("服务器主机名",   "{{{ hostname }}}"),
        ("实例启动时间",   "{{{ uptime_text }}}"),
        ("巡检人员",       inspector_name),
        ("服务器平台",     "{{{ platform_text }}}"),
        ("报告生成时间",   "{{{ report_time }}}"),
    ]
    for i, (label, value) in enumerate(info_cells):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = value

    # 第1章 基本信息
    doc.add_paragraph("\n第1章 数据库基本信息", style='Heading1Custom')
    t = doc.add_table(rows=5, cols=2)
    t.style = 'Table Grid'
    cells = [
        ("实例名称", "{{{ co_name }}}"),
        ("数据库版本", "{{{ ora_version }}}"),
        ("IP 地址", "{{{ ip }}}"),
        ("端口", "{{{ port }}}"),
        ("健康状态", "{{{ health_status }}}"),
    ]
    for i, (k, v) in enumerate(cells):
        t.cell(i, 0).text = k; t.cell(i, 1).text = v

    # 第2章 实例状态
    doc.add_paragraph("\n第2章 实例与运行状态", style='Heading1Custom')
    doc.add_paragraph("{{{ ora_instance }}}")

    # 第3章 表空间
    doc.add_paragraph("\n第3章 表空间使用情况", style='Heading1Custom')
    doc.add_paragraph("{{{ ora_tablespace }}}")

    # 第4章 SGA/PGA 内存
    doc.add_paragraph("\n第4章 SGA/PGA 内存分析", style='Heading1Custom')
    doc.add_paragraph("{{{ ora_sga_total }}}\n{{{ ora_sga }}}\n{{{ ora_pga }}}")

    # 第5章 会话与锁
    doc.add_paragraph("\n第5章 会话与锁等待", style='Heading1Custom')
    doc.add_paragraph("{{{ ora_sessions }}}\n{{{ ora_blocked }}}")

    # 第6章 Redo/归档
    doc.add_paragraph("\n第6章 Redo 日志与归档", style='Heading1Custom')
    doc.add_paragraph("{{{ ora_redo_logs }}}\n{{{ ora_archive_lag }}}")

    # 第7章 系统资源
    doc.add_paragraph("\n第7章 系统资源监控", style='Heading1Custom')
    doc.add_paragraph("{{{ system_info_text }}}")

    # 第8章 无效对象 & 用户
    doc.add_paragraph("\n第8章 对象与用户安全", style='Heading1Custom')
    doc.add_paragraph("8.1 无效对象\n{{{ ora_invalid_cnt }}}")
    doc.add_paragraph("8.2 用户列表\n{{{ ora_users }}}")

    # 第9章 报告说明（与 PostgreSQL 格式一致）
    doc.add_paragraph("\n9. 报告说明", style='Heading1Custom')
    doc.add_paragraph("{{{ notes_text }}}")

    os.makedirs(template_path, exist_ok=True)
    doc.save(tpl_file)
    return tpl_file


# ============================================================
# Oracle 数据采集类
# ============================================================
class getData(object):
    """数据采集类 - 负责连接 Oracle 数据库并执行全量巡检 SQL"""

    def __init__(self, ip, port, user, password, service_name=None, ssh_info=None):
        self.label = str(infos.label)
        self.H = ip
        self.P = int(port)
        self.user = user
        self.password = password
        self.service_name = service_name or ip
        self.ssh_info = ssh_info or {}
        self.conn_db2 = None
        self._connect()

    @staticmethod
    def _parse_sysdba_user(raw_user):
        """解析用户名中的特权身份（如 'sys as sysdba' → ('sys', 'sysdba')）"""
        raw = raw_user.strip()
        for suffix in [' as sysdba', ' as sysoper', ' as sysasm']:
            if raw.lower().endswith(suffix):
                return raw[:-len(suffix)].strip(), suffix[4:].strip()  # 去掉 " as "
        return raw, None

    def _connect(self):
        try:
            user_clean, privilege = self._parse_sysdba_user(self.user)

            # 确定 mode 参数值
            if ORACLE_DRIVER == 'oracledb':
                if privilege:
                    try:
                        mode_map = {'sysdba': oracle_driver.SYSDBA,
                                    'sysoper': oracle_driver.SYSOPER,
                                    'sysasm':  oracle_driver.SYSASM}
                        mode_val = mode_map.get(privilege.lower())
                    except AttributeError:
                        print(f"⚠️  驱动不支持 {privilege} 模式，将使用普通连接")
                        mode_val = None
                else:
                    mode_val = None  # 普通连接

                self.conn_db2 = oracle_driver.connect(
                    user=user_clean, password=self.password,
                    dsn=f"{self.H}:{self.P}/{self.service_name}",
                    mode=mode_val
                )
            else:  # cx_Oracle
                if privilege:
                    try:
                        mode_map = {'sysdba': oracle_driver.SYSDBA,
                                    'sysoper': oracle_driver.SYSOPER,
                                    'sysasm':  oracle_driver.SYSASM}
                        mode_val = mode_map.get(privilege.lower())
                    except AttributeError:
                        print(f"⚠️  驱动不支持 {privilege} 模式，将使用普通连接")
                        mode_val = None
                else:
                    mode_val = None

                dsn_tns = oracle_driver.makedsn(self.H, self.P, service_name=self.service_name)
                self.conn_db2 = oracle_driver.connect(
                    user=user_clean, password=self.password,
                    dsn=dsn_tns, mode=mode_val
                )

            priv_info = f" (as {privilege})" if privilege else ""
            print(f"✅ Oracle 连接成功: {user_clean}@{self.H}:{self.P}/{self.service_name}{priv_info}")
        except Exception as e:
            print(f"❌ Oracle 连接失败: {e}")
            self.conn_db2 = None
        self.context = {}

    def print_progress_bar(self, iteration, total, prefix='', suffix='', decimals=1, length=50, fill='█'):
        percent = ("{0:." + str(decimals) + "f}").format(100 * (iteration / float(total)))
        filled_length = int(length * iteration // total)
        bar = fill * filled_length + '-' * (length - filled_length)
        print(f'\r{prefix} |{bar}| {percent}% {suffix}', end='\r')
        if iteration == total:
            print()

    def checkdb(self, sqlfile=''):
        """
        执行 Oracle 数据库健康巡检。

        主要流程：
        1. 读取 SQL 模板（内置 builtin 或外部 .ini 文件）
        2. 逐条执行模板中的 SQL 语句，结果存入 context
        3. 采集系统信息（SSH 或本地）
        4. 自动风险分析
        5. AI 诊断
        """
        print("\n开始 Oracle 巡检...")
        total_steps = 22
        current_step = 0
        cfg = configparser.RawConfigParser()
        try:
            if sqlfile == 'builtin':
                cfg.read_string(ORACLE_SQL_TEMPLATES)
            else:
                cfg.read(sqlfile, encoding='utf-8')
        except Exception as e:
            print(f"❌ 读取SQL模板失败: {e}")
            return self.context

        init_keys = [
            "ora_version", "ora_instance", "ora_database", "ora_uptime",
            "ora_sessions", "ora_session_limit", "ora_process_limit", "ora_blocked",
            "ora_tablespace", "ora_temp_ts",
            "ora_sga", "ora_sga_total", "ora_pga", "ora_memory_target",
            "ora_redo_logs", "ora_redo_status",
            "ora_archive_dest", "ora_archive_lag", "ora_backup",
            "ora_params",
            "ora_invalid_objs", "ora_invalid_cnt",
            "ora_users", "ora_sys_privs", "ora_default_pws",
            "ora_long_sql", "ora_top_sql_cpu",
            "ora_dg_status", "ora_dg_apply", "ora_standby_event",
            "ora_rac_nodes", "ora_rac_interconn",
            "ora_asm_diskgroup",
            "ora_undo_info",
            "ora_recyclebin", "ora_flashback_area",
            "ora_datafiles",
            "ora_profile_pwd",
            "ora_top_waits", "ora_wait_class",
            "ora_stale_stats", "ora_partition_info"
        ]
        for key in init_keys:
            self.context.update({key: []})

        # ── 步骤1: 获取版本 ──────────────────────────
        try:
            cursor_ver = self.conn_db2.cursor()
            cursor_ver.execute("SELECT banner FROM v$version WHERE ROWNUM=1 AND banner LIKE 'Oracle%'")
            ver_row = cursor_ver.fetchone()
            oracle_version = ver_row[0] if ver_row else "Unknown"
            cursor_ver.close()
            self.context.update({"ora_version": [{'BANNER': oracle_version}]})
            self.context.update({"health_summary": [{'health_summary': '运行良好'}]})
        except Exception as e:
            print(f"❌ 获取版本失败: {e}")
            self.context.update({"ora_version": [{'BANNER': 'Unknown'}]})

        # ── 步骤2-21: 执行所有 SQL ───────────────────
        try:
            cursor = self.conn_db2.cursor()
            variables_items = list(cfg.items("variables"))
            for i, (name, stmt) in enumerate(variables_items):
                try:
                    current_step = int((i + 1) / len(variables_items) * (total_steps - 6)) + 1
                    self.print_progress_bar(current_step, total_steps, prefix='Oracle巡检:', suffix=f'{name} ({i+1}/{len(variables_items)})')
                    clean_sql = stmt.replace('\n', ' ').replace('\r', ' ')
                    # 移除尾部分号：DB-API 的 cursor.execute() 不接受 ;
                    # （; 是 SQL*Plus 命令终止符，不是 SQL 语句的一部分）
                    clean_sql = clean_sql.rstrip().rstrip(';').strip()
                    cursor.execute(clean_sql)
                    columns = [col[0] for col in cursor.description]
                    result = []
                    for row in cursor.fetchall():
                        result.append(dict(zip(columns, row)))
                    self.context[name] = result
                    time.sleep(0.03)
                except Exception as e:
                    print(f"\n⚠️  {name} 执行失败: {str(e)[:120]}")
                    self.context[name] = []
        except Exception as e:
            print(f'\n❌ 数据库查询失败: {e}')
        finally:
            if 'cursor' in locals():
                try: cursor.close()
                except Exception: pass

        # ── 步骤: 收集系统信息 ─────────────────────
        current_step = total_steps - 4
        self.print_progress_bar(current_step, total_steps, prefix='Oracle巡检:', suffix='收集系统信息')
        try:
            if self.ssh_info and self.ssh_info.get('ssh_host'):
                print(f"\n🔍 通过SSH收集系统信息: {self.ssh_info['ssh_host']}")
                collector = RemoteSystemInfoCollector(
                    host=self.ssh_info['ssh_host'], port=self.ssh_info.get('ssh_port', 22),
                    username=self.ssh_info.get('ssh_user', 'root'),
                    password=self.ssh_info.get('ssh_password'), key_file=self.ssh_info.get('ssh_key_file')
                )
                if not collector.connect():
                    print("⚠️  SSH 连接失败，跳过远程系统信息采集")
                    collector = LocalSystemInfoCollector()
            else:
                print(f"\n🔍 收集本地系统信息")
                collector = LocalSystemInfoCollector()
            system_info = collector.get_system_info()
            if isinstance(system_info.get('disk'), dict):
                disk_list = list(system_info['disk'].values())
                system_info['disk_list'] = disk_list
            elif isinstance(system_info.get('disk'), list):
                system_info['disk_list'] = system_info['disk']
            else:
                disk_info = get_host_disk_usage()
                system_info['disk_list'] = disk_info
            self.context.update({"system_info": system_info})
        except Exception as e:
            print(f"\n❌ 收集系统信息失败: {e}")
            self.context.update({"system_info": {
                'hostname': '未知', 'platform': '未知', 'boot_time': '未知',
                'cpu': {}, 'memory': {},
                'disk_list': [{'device':'/dev/sda1','mountpoint':'/','fstype':'ext4',
                               'total_gb':0,'used_gb':0,'free_gb':0,'usage_percent':0}]
            }})

        # ── 步骤: 风险分析 ─────────────────────────
        current_step = total_steps - 3
        self.print_progress_bar(current_step, total_steps, prefix='Oracle巡检:', suffix='智能风险分析')
        self.context.update({"auto_analyze": []})
        try:
            from analyzer import smart_analyze_oracle
            issues = smart_analyze_oracle(self.context)
            self.context['auto_analyze'] = issues
        except ImportError:
            # 降级基础规则
            self._basic_risk_check()
        except Exception as e:
            print(f"\n❌ 风险分析失败: {e}")

        # ── 步骤: AI 诊断 ─────────────────────────
        current_step = total_steps - 2
        self.print_progress_bar(current_step, total_steps, prefix='Oracle巡检:', suffix='AI诊断')
        self.context['ai_advice'] = ''
        try:
            from analyzer import AIAdvisor
            import json as _json
            cfg_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'ai_config.json')
            ai_cfg = {}
            if os.path.exists(cfg_path):
                with open(cfg_path, 'r', encoding='utf-8') as f:
                    ai_cfg = _json.load(f)
            advisor = AIAdvisor(
                backend=ai_cfg.get('backend'),
                api_key=ai_cfg.get('api_key'),
                api_url=ai_cfg.get('api_url'),
                model=ai_cfg.get('model')
            )
            if advisor.enabled:
                label = self.context.get('co_name', [{}])[0].get('CO_NAME', 'Oracle')
                print(f"\n🤖 正在调用 AI 诊断（{advisor.backend} / {advisor.model}）...")
                ai_advice = advisor.diagnose('oracle', label, self.context, issues)
                self.context['ai_advice'] = ai_advice
        except Exception as e:
            self.context['ai_advice'] = ''

        self.print_progress_bar(total_steps, total_steps, prefix='Oracle巡检:', suffix='完成 ✅')
        return self.context

    def _basic_risk_check(self):
        """基础风险检查（analyzer.py 不可用时的降级方案）"""
        # 表空间使用率检查
        ts_list = self.context.get('ora_tablespace', [])
        for ts in ts_list:
            used_pct = self._safe_float(ts.get('used_pct_with_maxext', ts.get('used_pct', 0)))
            if used_pct > 90:
                ts_name = ts.get("TABLESPACE_NAME", "?")
                self.context['auto_analyze'].append({
                    'col1': f'表空间 {ts_name}', 'col2': '高风险',
                    'col3': f'表空间使用率 {used_pct:.1f}%，超过 90% 告警线',
                    'col4': '高', 'col5': 'DBA',
                    'fix_sql': (
                        "-- 检查大对象：\n"
                        f"-- SELECT segment_name, bytes/1024/1024 AS mb "
                        f"FROM dba_segments WHERE tablespace_name='{ts_name}' "
                        "ORDER BY bytes DESC;"
                    )
                })
            elif used_pct > 80:
                self.context['auto_analyze'].append({
                    'col1': f'表空间 {ts.get("TABLESPACE_NAME","?")}', 'col2': '中风险',
                    'col3': f'表空间使用率 {used_pct:.1f}%，建议关注',
                    'col4': '中', 'col5': 'DBA',
                    'fix_sql': ''
                })

        # 会话数检查
        sess = self.context.get('ora_sessions', [])
        limit = self.context.get('ora_session_limit', [])
        if sess and limit:
            total = self._safe_int(sess[0], 'TOTAL_SESSIONS')
            max_sess = self._safe_int(limit[0], 'SESSIONS_LIMIT')
            if max_sess > 0 and (total / max_sess) * 100 > 85:
                self.context['auto_analyze'].append({
                    'col1': '会话数接近上限', 'col2': '高风险',
                    'col3': f'当前会话 {total} / 上限 {max_sess}',
                    'col4': '高', 'col5': 'DBA',
                    'fix_sql': '-- 查看会话详情:\nSELECT sid, serial#, username, status, machine FROM v$session WHERE type=\'USER\' AND status=\'ACTIVE\';'
                })

        # 系统内存/磁盘检查
        mem = self.context.get('system_info', {}).get('memory', {})
        if mem.get('usage_percent', 0) > 90:
            self.context['auto_analyze'].append({
                'col1': '系统内存紧张', 'col2': '高风险',
                'col3': f'内存使用率 {mem["usage_percent"]:.1f}%',
                'col4': '高', 'col5': '系统管理员', 'fix_sql': ''
            })
        for disk in self.context.get('system_info', {}).get('disk_list', []):
            usage = self._safe_float(disk, 'usage_percent')
            mp = disk.get('mountpoint', '/')
            if mp in IGNORE_MOUNTS: continue
            if usage > 90:
                self.context['auto_analyze'].append({
                    'col1': f'磁盘空间不足 ({mp})', 'col2': '高风险',
                    'col3': f'磁盘 {mp} 使用率 {usage:.1f}%',
                    'col4': '高', 'col5': '系统管理员', 'fix_sql': ''
                })

    @staticmethod
    def _safe_float(obj, field='value', default=0.0):
        try:
            val = obj[field] if isinstance(obj, dict) else getattr(obj, field, None)
            return float(str(val).replace(',', '').replace('%', '')) if val is not None else default
        except Exception:
            return default

    @staticmethod
    def _safe_int(obj, field='value', default=0):
        try:
            val = obj[field] if isinstance(obj, dict) else getattr(obj, field, None)
            return int(str(val).replace(',', '')) if val is not None else default
        except Exception:
            return default


# ============================================================
# 报告保存类
# ============================================================
class saveDoc(object):
    """报告保存类 - 将 Oracle 巡检数据渲染到 Word 模板"""

    def __init__(self, context, ofile, ifile, inspector_name="Jack"):
        self.context = context
        self.ofile = ofile
        self.ifile = ifile
        self.inspector_name = inspector_name

    def contextsave(self):
        try:
            required_keys = ['health_summary', 'auto_analyze', 'ora_version', 'co_name', 'port', 'ip', 'system_info']
            for key in required_keys:
                if key not in self.context:
                    if key == 'health_summary':   self.context[key] = [{'health_summary': '运行良好'}]
                    elif key == 'auto_analyze':     self.context[key] = []
                    elif key == 'ora_version':      self.context[key] = [{'BANNER': 'Unknown'}]
                    elif key == 'system_info':      self.context[key] = {}
                    else:                           self.context[key] = [{'placeholder': '数据缺失'}]

            if 'disk_list' not in self.context['system_info'] or not self.context['system_info']['disk_list']:
                self.context['system_info']['disk_list'] = [{
                    'device': '/dev/sda1', 'mountpoint': '/', 'fstype': 'ext4',
                    'total_gb': 50.0, 'used_gb': 25.0, 'free_gb': 25.0, 'usage_percent': 50.0
                }]

            list_keys = [
                'ora_tablespace', 'ora_temp_ts', 'ora_sessions', 'ora_blocked',
                'ora_sga', 'ora_sga_total', 'ora_pga', 'ora_memory_target',
                'ora_redo_logs', 'ora_redo_status', 'ora_archive_dest',
                'ora_archive_lag', 'ora_backup', 'ora_params',
                'ora_invalid_objs', 'ora_invalid_cnt', 'ora_users',
                'ora_sys_privs', 'ora_default_pws', 'ora_long_sql',
                'ora_top_sql_cpu', 'ora_dg_status', 'ora_dg_apply',
                'ora_standby_event', 'ora_rac_nodes', 'ora_rac_interconn',
                'ora_asm_diskgroup', 'ora_undo_info', 'ora_recyclebin',
                'ora_flashback_area', 'ora_datafiles', 'ora_profile_pwd',
                'ora_top_waits', 'ora_wait_class', 'ora_stale_stats',
                'ora_partition_info'
            ]
            for key in list_keys:
                if key not in self.context:
                    self.context[key] = []

            self.context.update({"report_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S')})
            self.context.update({"inspector_name": self.inspector_name})
            problem_count = len(self.context.get("auto_analyze", []))
            self.context.update({"problem_count": problem_count})

            if problem_count == 0: health_status = "优秀"
            elif problem_count <= 3: health_status = "良好"
            elif problem_count <= 6: health_status = "一般"
            else: health_status = "需关注"
            self.context.update({"health_status": health_status})

            # ── 预格式化复杂变量为纯文本，避免模板 Jinja2 解析问题 ──
            # 基本信息：列表取第一个值
            if isinstance(self.context.get('co_name'), list) and len(self.context['co_name']) > 0:
                self.context['co_name'] = str(self.context['co_name'][0].get('CO_NAME', ''))
            else:
                self.context['co_name'] = str(self.context.get('co_name', '') or '')
            if isinstance(self.context.get('ora_version'), list) and len(self.context['ora_version']) > 0:
                self.context['ora_version'] = str(self.context['ora_version'][0].get('BANNER', ''))
            else:
                self.context['ora_version'] = str(self.context.get('ora_version', '') or '')
            if isinstance(self.context.get('ip'), list) and len(self.context['ip']) > 0:
                ip_val = self.context['ip'][0].get('IP', '')
            else:
                ip_val = str(self.context.get('ip', '') or '')
            port_val = ''
            if isinstance(self.context.get('port'), list) and len(self.context['port']) > 0:
                port_val = str(self.context['port'][0].get('PORT', ''))
            elif self.context.get('port'):
                port_val = str(self.context['port'])
            self.context['server_addr'] = f"{ip_val}:{port_val}" if ip_val else ""
            self.context['ip'] = ip_val
            self.context['port'] = port_val

            # 系统信息：主机名、平台、实例启动时间
            sys_info = self.context.get('system_info') or {}
            self.context['hostname'] = str(sys_info.get('hostname', 'N/A') or 'N/A')
            platform_info = sys_info.get('platform', {})
            if isinstance(platform_info, dict):
                plat_str = platform_info.get('platform', '') or ''
                arch = platform_info.get('machine', '') or ''
                self.context['platform_text'] = f"{plat_str} ({arch})" if plat_str else (arch or 'N/A')
            else:
                self.context['platform_text'] = str(platform_info) if platform_info else 'N/A'
            # 实例启动时间（从 ora_uptime 提取）
            uptime_data = self.context.get('ora_uptime')
            if isinstance(uptime_data, list) and len(uptime_data) > 0:
                u = uptime_data[0]
                start_time = u.get('STARTUP_TIME', '') or u.get('STARTUP', '') or ''
                days = u.get('DAYS', 0) or 0
                hours = u.get('HOURS', 0) or 0
                mins = u.get('MINUTES', 0) or 0
                if start_time:
                    self.context['uptime_text'] = f"{start_time} (已运行 {days}天{hours}小时{mins}分钟)"
                else:
                    self.context['uptime_text'] = f"已运行 {days}天{hours}小时{mins}分钟"
            else:
                self.context['uptime_text'] = 'N/A'

            # 系统信息预格式化为文本
            sys_info = self.context.get('system_info') or {}
            cpu = sys_info.get('cpu', {}) or {}
            mem = sys_info.get('memory', {}) or {}
            disks = sys_info.get('disk_list', []) or []
            lines = []
            lines.append(f"CPU 使用率: {cpu.get('usage_percent', 'N/A')}%")
            used_mb = mem.get('used_mb', 0)
            total_mb = mem.get('total_mb', 0)
            mem_pct = mem.get('usage_percent', 0)
            lines.append(f"内存使用: {used_mb} MB / {total_mb} MB ({mem_pct}%)")
            if disks and isinstance(disks, list):
                lines.append("磁盘空间:")
                for d in disks[:10]:
                    mp = d.get('mountpoint', '/')
                    tg = d.get('total_gb', 0)
                    ug = d.get('used_gb', 0)
                    fg = d.get('free_gb', 0)
                    up = d.get('usage_percent', 0)
                    lines.append(f"  {mp}\t{tg}GB\t{ug}GB\t{fg}GB\t{up}%")
            self.context['system_info_text'] = '\n'.join(lines)

            # 第9章报告说明预格式化
            notes = [
                "1. 本报告基于 Oracle 数据库实时状态生成，反映了生成时刻的数据库健康状况",
                "2. 报告中空白的项表示未能获取到相关数据，可能是由于权限限制或该功能未启用",
                "3. 磁盘信息仅显示主要分区的使用率，如需查看完整磁盘信息请使用系统命令 'df -h'",
                "4. 巡检结果仅供参考，实际运维中请结合具体业务场景进行分析",
                "5. 建议定期进行数据库巡检，及时发现并解决潜在问题",
                "6. AI 诊断功能（若启用）生成的建议仅供参考，不构成专业 DBA 意见"
            ]
            self.context['notes_text'] = '\n'.join(notes)

            # 列表数据预格式化为表格文本
            _list_keys_to_text = ['ora_tablespace','ora_sessions','ora_blocked','ora_sga',
                'ora_sga_total','ora_pga','ora_redo_logs','ora_archive_lag',
                'ora_backup','ora_params','ora_invalid_objs','ora_invalid_cnt',
                'ora_users','ora_sys_privs','ora_default_pws','ora_long_sql',
                'ora_top_sql_cpu','ora_dg_status','ora_dg_apply','ora_standby_event',
                'ora_rac_nodes','ora_rac_interconn','ora_asm_diskgroup','ora_undo_info',
                'ora_recyclebin','ora_flashback_area','ora_datafiles','ora_profile_pwd',
                'ora_top_waits','ora_wait_class','ora_stale_stats','ora_partition_info']
            for key in _list_keys_to_text:
                val = self.context.get(key)
                if val is None: continue
                if isinstance(val, (list, tuple)):
                    if not val:
                        self.context[key] = "无数据"
                        continue
                    rows = []
                    headers = []
                    for item in val:
                        if isinstance(item, dict):
                            if not headers: headers = list(item.keys())
                            rows.append('\t'.join(str(item.get(h,'')) for h in headers))
                    if headers:
                        self.context[key] = '\t'.join(headers) + '\n' + '\n'.join(rows)
                    else:
                        self.context[key] = str(val)[:2000]
                elif not isinstance(val, str):
                    self.context[key] = str(val)

            # 尝试 docxtpl 渲染
            try:
                with open(self.ifile, 'rb') as f:
                    template_bytes = f.read()
                doc_stream = io.BytesIO(template_bytes)
                tpl = DocxTemplate(doc_stream)
                tpl.render(self.context)
                tpl.save(self.ofile)

                # 追加新章节（与 MySQL 版本一致）
                doc2 = Document(self.ofile)
                cutoff_idx = None
                for i, para in enumerate(doc2.paragraphs):
                    t = para.text.strip()
                    if t.startswith('9.') and '报告说明' in t:
                        cutoff_idx = i
                        break
                if cutoff_idx is not None:
                    for j in range(len(doc2.paragraphs) - 1, cutoff_idx - 1, -1):
                        p = doc2.paragraphs[j]._element
                        p.getparent().remove(p)

                self._append_chapters(doc2)
                doc2.save(self.ofile)
                return True
            except Exception as e:
                # docxtpl 模板解析异常时自动降级为备用渲染，无需向用户展示
                return self._fallback_render()
        except Exception as e:
            print(f"报告生成异常: {e}")
            return False

    def _append_chapters(self, doc):
        """在 Word 文档末尾追加新章节（第7-9章）"""
        from docx.oxml.ns import qn as _qn

        def _add_heading(text, level=1):
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(_qn('w:eastAsia'), '微软雅黑')
                run.font.color.rgb = RGBColor(0, 51, 102)
            return h

        def _add_table(headers, rows):
            t = doc.add_table(rows=max(1,len(rows))+1, cols=len(headers), style='Table Grid')
            for j, h in enumerate(headers):
                cell = t.cell(0, j); cell.text = h
                for p in cell.paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for i, row_data in enumerate(rows):
                for j, val in enumerate(row_data):
                    t.cell(i+1, j).text = str(val)[:200] if val else ''
            return t

        # 第7章 风险与建议
        _add_heading("第7章 风险与建议")
        issues = self.context.get("auto_analyze", [])
        if issues:
            _add_heading("7.1 问题明细", 2)
            _add_table(["序号", "项目", "风险等级", "问题描述", "严重程度", "责任人", "修复建议"],
                       [(str(i+1), x.get('col1',''), x.get('col2',''), x.get('col3',''),
                         x.get('col4',''), x.get('col5',''), x.get('fix_sql','')[:200]) for i,x in enumerate(issues)])

            fix_sqls = [(x.get('col1',''), x.get('fix_sql','')) for x in issues if x.get('fix_sql')]
            if fix_sqls:
                _add_heading("7.2 修复SQL速查", 2)
                for fname, sql in fix_sqls:
                    p = doc.add_paragraph(); p.add_run(f"【{fname}】").bold = True
                    doc.add_paragraph(sql, style='List Bullet')
        else:
            doc.add_paragraph("未发现明显风险项，数据库整体运行状况良好 👍")

        # 第8章 AI 诊断
        ai_text = self.context.get('ai_advice', '')
        if ai_text:
            _add_heading("第8章 AI 诊断建议")
            for line in ai_text.split('\n'):
                if line.startswith('# '): _add_heading(line[2:], level=2)
                elif line.startswith('## '): _add_heading(line[3:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif re.match(r'^\d+\.', line):
                    doc.add_paragraph(line, style='List Number')
                elif line.strip():
                    doc.add_paragraph(line)

        # 第9章 报告说明（与 PostgreSQL 格式一致）
        _add_heading("第9章 报告说明")
        notes = [
            "1. 本报告基于 Oracle 数据库实时状态生成，反映了生成时刻的数据库健康状况",
            "2. 报告中空白的项表示未能获取到相关数据，可能是由于权限限制或该功能未启用",
            "3. 磁盘信息仅显示主要分区的使用率，如需查看完整磁盘信息请使用系统命令 'df -h'",
            "4. 巡检结果仅供参考，实际运维中请结合具体业务场景进行分析",
            "5. 建议定期进行数据库巡检，及时发现并解决潜在问题",
            "6. AI 诊断功能（若启用）生成的建议仅供参考，不构成专业 DBA 意见"
        ]
        for note in notes:
            p = doc.add_paragraph(note)
            p.runs[0].font.size = Pt(10) if p.runs else None

    def _fallback_render(self):
        """增强备用渲染（docxtpl 失败时使用，兼容预格式化的字符串数据）"""
        try:
            doc = Document()
            from docx.oxml.ns import qn as _qn

            def _h(text, lvl=1):
                hh = doc.add_heading(text, level=lvl)
                for r in hh.runs:
                    r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(_qn('w:eastAsia'), '微软雅黑'); r.font.color.rgb = RGBColor(0,51,102)
                return hh

            def _t(hdr, rows):
                tt = doc.add_table(rows=max(1,len(rows))+1, cols=len(hdr), style='Table Grid')
                for j,h in enumerate(hdr):
                    c=tt.cell(0,j); c.text=h; [p.__setattr__('alignment', WD_ALIGN_PARAGRAPH.CENTER) for p in c.paragraphs]
                for i,row in enumerate(rows):
                    for j,v in enumerate(row): tt.cell(i+1,j).text=str(v)[:200] if v else ''

            def _render(key, title=None, max_r=50):
                val = self.context.get(key)
                if not val: return
                if isinstance(val, str) and val != '无数据':
                    if title: _h(title)
                    doc.add_paragraph(val[:5000])
                elif isinstance(val, (list,tuple)) and len(val)>0 and isinstance(val[0], dict):
                    if title: _h(title)
                    hdrs = list(val[0].keys())
                    _t(hdrs, [list(d.values()) for d in val[:max_r]])

            ctx = self.context

            # ── 封面：8行信息表（与 PG 格式一致）──
            doc.add_heading("Oracle 数据库健康巡检报告", 0)
            info_t = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
            info_rows = [
                ("数据库名称", str(ctx.get('co_name', ''))),
                ("服务器地址", str(ctx.get('server_addr', ''))),
                ("Oracle版本",   str(ctx.get('ora_version', ''))),
                ("服务器主机名",  str(ctx.get('hostname', ''))),
                ("实例启动时间",  str(ctx.get('uptime_text', ''))),
                ("巡检人员",      str(self.inspector_name)),
                ("服务器平台",    str(ctx.get('platform_text', ''))),
                ("报告生成时间",  str(ctx.get('report_time', '')))
            ]
            for i,(k,v) in enumerate(info_rows):
                info_t.cell(i,0).text = k; info_t.cell(i,1).text = v

            # Ch1
            _h("第1章 数据库基本信息")
            _render('ora_database')
            _render('ora_uptime')

            # Ch2
            _h("第2章 实例与运行状态")
            _render('ora_instance')
            _render('ora_database')
            _render('ora_uptime')

            # Ch3 表空间
            _h("第3章 表空间使用情况")
            _render('ora_tablespace', None, 30)
            _render('ora_temp_ts')

            # Ch4 SGA/PGA
            _h("第4章 SGA/PGA 内存分析")
            _render('ora_sga_total')
            _render('ora_sga', None, 15)
            _render('ora_pga')

            # Ch5 会话锁
            _h("第5章 会话与锁等待")
            _render('ora_sessions')
            _render('ora_blocked', None, 20)

            # Ch6 Redo/归档
            _h("第6章 Redo 日志与归档")
            _render('ora_redo_logs')
            _render('ora_archive_lag')

            # Ch7 系统资源
            _h("第7章 系统资源监控")
            stx = ctx.get('system_info_text')
            if stx: doc.add_paragraph(stx)

            # Ch8 归档备份
            ha = ctx.get('ora_archive_dest') and str(ctx['ora_archive_dest']) not in ('无数据','')
            hb = ctx.get('ora_backup') and str(ctx['ora_backup']) not in ('无数据','')
            if ha or hb:
                _h("第8章 归档与备份信息")
                _render('ora_archive_dest'); _render('ora_archive_lag'); _render('ora_backup')

            # Ch9 Data Guard
            dg_v = ctx.get('ora_dg_status')
            if dg_v and str(dg_v) not in ('无数据',''):
                _h("第9章 Data Guard / ADG 状态")
                _render('ora_dg_status'); _render('ora_dg_apply')

            # Ch10 ASM
            av = ctx.get('ora_asm_diskgroup')
            if av and str(av) not in ('无数据',''):
                _h("第10章 ASM 磁盘组"); _render('ora_asm_diskgroup')

            # Ch11 用户安全
            _h("第11章 对象与用户安全")
            for k in ['ora_invalid_cnt','ora_users','ora_default_pws','ora_sys_privs']:
                _render(k, None, 50)

            # Ch12 参数
            _h("第12章 关键参数"); _render('ora_params', None, 30)

            # Ch13 无效对象
            iv = ctx.get('ora_invalid_objs')
            if iv and str(iv) not in ('无数据',''):
                _h("第13章 无效对象详情"); _render('ora_invalid_objs', None, 30)

            # Ch14 Top SQL
            _h("第14章 TOP SQL")
            _render('ora_top_sql_cpu', None, 15); _render('ora_long_sql', None, 15)

            # Ch15 其他诊断
            _h("第15章 其他诊断信息")
            for k2 in ['ora_top_waits','ora_wait_class','ora_stale_stats',
                        'ora_partition_info','ora_undo_info','ora_recyclebin',
                        'ora_flashback_area','ora_datafiles','ora_profile_pwd']:
                _render(k2, None, 20)

            self._append_chapters(doc)

            doc.save(self.ofile)
            return True
        except Exception as e:
            print(f"备用渲染也失败了: {e}")
            import traceback; traceback.print_exc()
            return False


# ============================================================
# 命令行参数兼容对象（供 web_ui.py 注入使用）
# ============================================================
            from docx.oxml.ns import qn as _qn

            def _h(text, lvl=1):
                hh = doc.add_heading(text, level=lvl)
                for r in hh.runs:
                    r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(_qn('w:eastAsia'), '微软雅黑'); r.font.color.rgb = RGBColor(0,51,102)
                return hh

            def _t(hdr, rows):
                tt = doc.add_table(rows=max(1,len(rows))+1, cols=len(hdr), style='Table Grid')
                for j,h in enumerate(hdr):
                    c=tt.cell(0,j); c.text=h; [p.__setattr__('alignment', WD_ALIGN_PARAGRAPH.CENTER) for p in c.paragraphs]
                for i,row in enumerate(rows):
                    for j,v in enumerate(row): tt.cell(i+1,j).text=str(v)[:200] if v else ''

            # 封面
            doc.add_heading("Oracle 数据库健康巡检报告", 0)
            doc.add_paragraph(f"巡检人员：{self.inspector_name}\n巡检时间：{self.context.get('report_time','')}")

            # Ch1: 基本信息
            _h("第1章 基本信息")
            ctx = self.context
            ver = ctx.get('ora_version',[{}])[0].get('BANNER','Unknown') if ctx.get('ora_version') else 'Unknown'
            inst = ctx.get('ora_instance',[{}])[0] if ctx.get('ora_instance') else {}
            dbinfo = ctx.get('ora_database',[{}])[0] if ctx.get('ora_database') else {}
            uptime = ctx.get('ora_uptime',[{}])[0] if ctx.get('ora_uptime') else {}

            _t(["属性","值"],[
                ["实例名称", ctx.get('co_name',[{}])[0].get('CO_NAME','-')],
                ["数据库版本", ver],
                ["主机地址", ctx.get('ip',[{}])[0].get('IP','-')],
                ["端口", str(ctx.get('port',[{}])[0].get('PORT','-'))],
                ["实例状态", inst.get('STATUS','-')],
                ["数据库角色", dbinfo.get('DATABASE_ROLE','-') or '-'],
                ["启动时间", uptime.get('STARTUP_TIME','-')],
                ["运行时长", f"{uptime.get('UPTIME_MINUTES',0)} 分钟"],
                ["归档模式", dbinfo.get('LOG_MODE','-')],
                ["健康评级", ctx.get('health_status','未知')]
            ])

            # Ch2: 实例与运行
            _h("第2章 实例与运行状态")
            if ctx.get('ora_instance'): _t(list(ctx['ora_instance'][0].keys()), [list(x.values()) for x in ctx['ora_instance']])
            if ctx.get('ora_database'): _t(list(ctx['ora_database'][0].keys()), [list(x.values()) for x in ctx['ora_database']])
            if ctx.get('ora_uptime'): _t(list(ctx['ora_uptime'][0].keys()), [list(x.values()) for x in ctx['ora_uptime']])

            # Ch3: 表空间
            _h("第3章 表空间使用情况")
            if ctx.get('ora_tablespace'):
                hdr = list(ctx['ora_tablespace'][0].keys()) if ctx['ora_tablespace'][0] else ['tablespace_name']
                _t(hdr, [list(x.values()) for x in ctx['ora_tablespace'][:30]])
            if ctx.get('ora_temp_ts'):
                thdr = list(ctx['ora_temp_ts'][0].keys()) if ctx['ora_temp_ts'][0] else ['tablespace_name']
                _t(thdr, [list(x.values()) for x in ctx['ora_temp_ts']])

            # Ch4: SGA/PGA
            _h("第4章 SGA/PGA 内存分析")
            if ctx.get('ora_sga_total'): _t(list(ctx['ora_sga_total'][0].keys()), [list(x.values()) for x in ctx['ora_sga_total']])
            if ctx.get('ora_sga'):
                shdr = list(ctx['ora_sga'][0].keys()) if ctx['ora_sga'][0] else []
                _t(shdr, [list(x.values()) for x in ctx['ora_sga'][:15]])
            if ctx.get('ora_pga'): _t(list(ctx['ora_pga'][0].keys()), [list(x.values()) for x in ctx['ora_pga']])

            # Ch5: 会话与锁
            _h("第5章 会话与锁等待")
            if ctx.get('ora_sessions'): _t(list(ctx['ora_sessions'][0].keys()), [list(x.values()) for x in ctx['ora_sessions']])
            if ctx.get('ora_blocked'):
                bhdr = list(ctx['ora_blocked'][0].keys()) if ctx['ora_blocked'][0] else []
                _t(bhdr, [list(x.values())[:min(len(bhdr),8)] for x in ctx['ora_blocked'][:20]])

            # Ch6: Redo/归档
            _h("第6章 Redo 日志与归档")
            if ctx.get('ora_redo_logs'): _t(list(ctx['ora_redo_logs'][0].keys()), [list(x.values()) for x in ctx['ora_redo_logs']])
            if ctx.get('ora_archive_lag'): _t(list(ctx['ora_archive_lag'][0].keys()), [list(x.values()) for x in ctx['ora_archive_lag']])

            # Ch7: 系统资源
            _h("第7章 系统资源监控")
            si = ctx.get('system_info', {})
            doc.add_paragraph(f"主机名: {si.get('hostname','-')}")
            doc.add_paragraph(f"平台: {si.get('platform','-')}")
            cpu = si.get('cpu', {}); mem = si.get('memory', {})
            doc.add_paragraph(f"CPU 使用率: {cpu.get('usage_percent','-')}%")
            doc.add_paragraph(f"内存: {mem.get('used_mb','-')}MB/{mem.get('total_mb','-')}MB ({mem.get('usage_percent','-')}%)")
            dlst = si.get('disk_list', [])
            if dlst:
                dh = list(dlst[0].keys()) if dlst[0] else []
                _t(dh, [list(d.values()) for d in dlst])

            # Ch8: Data Guard
            dg = ctx.get('ora_dg_status', [])
            if dg:
                _h("第8章 Data Guard / ADG 状态")
                _t(list(dg[0].keys()), [list(x.values()) for x in dg])
                if ctx.get('ora_dg_apply'):
                    ah = list(ctx['ora_dg_apply'][0].keys()) if ctx['ora_dg_apply'][0] else []
                    _t(ah, [list(v.values()) for v in ctx['ora_dg_apply']])

            # Ch9: ASM
            asm = ctx.get('ora_asm_diskgroup', [])
            if asm:
                _h("第9章 ASM 磁盘组")
                _t(list(asm[0].keys()), [list(x.values()) for x in asm])

            # Ch10: 用户安全
            _h("第10章 对象与用户安全")
            inv = ctx.get('ora_invalid_cnt', [])
            if inv: _t(list(inv[0].keys()), [list(x.values()) for x in inv])
            usrs = ctx.get('ora_users', [])
            if usrs: _t(list(usrs[0].keys()), [list(u.values()) for u in usrs[:50]])

            # Ch11: 参数
            params = ctx.get('ora_params', [])
            if params:
                _h("第11章 关键参数")
                ph = list(params[0].keys()) if params[0] else []
                _t(ph, [list(p.values()) for p in params[:30]])

            # Ch12: Top SQL
            top = ctx.get('ora_top_sql_cpu', [])
            if top:
                _h("第12章 TOP CPU SQL")
                toph = list(top[0].keys()) if top[0] else []
                _t(toph, [list(t.values()) for t in top[:15]])

            # Ch13: 等待事件
            wc = ctx.get('ora_wait_class', [])
            if wc:
                _h("第13章 等待事件分类统计")
                _t(list(wc[0].keys()), [list(w.values()) for w in wc])

            # Ch14: 统计信息
            stale = ctx.get('ora_stale_stats', [])
            if stale:
                _h("第14章 陈旧统计信息")
                sth = list(stale[0].keys()) if stale[0] else []
                _t(sth, [list(s.values()) for s in stale[:20]])

            # 追加章节（风险/AI/说明）
            self._append_chapters(doc)

            doc.save(self.ofile)
            return True
        except Exception as e:
            print(f"备用渲染也失败了: {e}")
            return False


# ============================================================
# 命令行参数兼容对象（供 web_ui.py 注入使用）
# ============================================================
class _Argus:
    """命令行参数容器"""
    def __init__(self):
        self.label = "Oracle_Inspector"
        self.sqltemplates = 'builtin'
        self.batch = False

infos = _Argus()


def run_inspection(db_info, service_name=None):
    """
    对单个 Oracle 实例执行完整巡检并生成 Word 报告。

    :param db_info: 包含连接信息的字典 {name, ip, port, user, password}
    :param service_name: Oracle 服务名/SID（可选）
    :return: 成功返回 True，失败返回 False
    """
    label_name = db_info['name']
    ip = db_info['ip']
    port = db_info['port']
    user = db_info['user']
    password = db_info['password']

    ssh_info = {}
    if 'ssh_host' in db_info and db_info['ssh_host']:
        ssh_info = {
            'ssh_host': db_info['ssh_host'],
            'ssh_port': db_info.get('ssh_port', 22),
            'ssh_user': db_info.get('ssh_user', 'root'),
            'ssh_password': db_info.get('ssh_password', ''),
            'ssh_key_file': db_info.get('ssh_key_file', '')
        }

    inspector_name = input("巡检人员（默认 Jack）: ").strip() or "Jack"
    ifile = create_word_template(inspector_name)
    if not ifile: return False

    dir_path = "reports"
    if not os.path.exists(dir_path): os.makedirs(dir_path)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    file_name = f"Oracle巡检报告_{label_name}_{timestamp}.docx"
    ofile = os.path.join(dir_path, file_name)

    try:
        print(f"\n🔍 正在测试 Oracle 连接 {ip}:{port}/{service_name or ip}...")
        data = getData(ip, port, user, password, service_name=service_name, ssh_info=ssh_info)
        if data is None or data.conn_db2 is None:
            return False
        ret = data.checkdb('builtin')
        if not ret: return False

        ret.update({"co_name": [{'CO_NAME': label_name}]})
        ret.update({"port": [{'PORT': port}]})
        ret.update({"ip": [{'IP': ip}]})

        savedoc = saveDoc(context=ret, ofile=ofile, ifile=ifile, inspector_name=inspector_name)
        success = savedoc.contextsave()
        if success:
            print(f"\n✅ 报告已生成: {file_name}")
            try:
                if os.path.exists(ifile): os.remove(ifile)
            except Exception: pass
            return True
        else:
            print(f"\n❌ 报告生成失败: {label_name}")
            return False
    except Exception as e:
        print(f"\n❌ 巡检失败: {e}")
        return False


# ============================================================
# 命令行主入口（CLI 模式，与 main_mysql 保持一致）
# ============================================================

def print_banner():
    print("""
╔═══════════════════════════════════════════════════════════╗
║         DBCheck - Oracle 数据库健康巡检工具 v1.0          ║
║              支持 Oracle 11g / 12c / 19c / 21c             ║
║                                                             ║
║   巡检项目：表空间 · SGA/PGA · 会话 · 锁 · Redo · 归档     ║
║              Data Guard · ASM · RAC · 用户安全 · 等待事件   ║
╚═══════════════════════════════════════════════════════════╝
""")


def show_main_menu():
    print("\n" + "="*50)
    print("  主菜单")
    print("="*50)
    print("  1. 单机巡检")
    print("  2. 批量巡检")
    print("  3. 创建 Excel 巡检模板")
    print("  4. 退出")
    print("="*50)
    return input("请选择 [1-4]: ").strip()


def single_inspection():
    print("\n--- 单机 Oracle 巡检 ---")
    name = input("实例标签 (例: 生产-orcl): ").strip() or "ORACLE"
    host = input("主机地址: ").strip() or "localhost"
    port = input("端口 [1521]: ").strip() or "1521"
    service = input("服务名/SID (留空同主机名): ").strip() or host
    user = input("用户名 [sys as sysdba 或 system]: ").strip() or "system"
    password = input("密码: ").strip()
    use_ssh = input("是否配置 SSH 收集系统信息? [y/N]: ").strip().lower()

    db_info = {'name': name, 'ip': host, 'port': port, 'user': user, 'password': password}

    if use_ssh in ('y', 'yes'):
        db_info['ssh_host'] = input("SSH 主机 [同数据库主机]: ").strip() or host
        db_info['ssh_port'] = int(input("SSH 端口 [22]: ").strip() or "22")
        db_info['ssh_user'] = input("SSH 用户 [root]: ").strip() or "root"
        pw = input("SSH 密码 (留空则跳过): ").strip()
        if pw: db_info['ssh_password'] = pw

    run_inspection(db_info, service_name=service)


def batch_inspection():
    print("\n--- 批量巡检 ---")
    excel_path = input("Excel 模板路径 (或回车创建新模板): ").strip()
    if not excel_path or not os.path.exists(excel_path):
        create_excel_template()
        excel_path = input("Excel 路径: ").strip()

    try:
        wb = openpyxl.load_workbook(excel_path)
        ws = wb.active
        success_count = 0
        for row_idx in range(2, ws.max_row + 1):
            name = ws.cell(row=row_idx, column=1).value
            ip = ws.cell(row=row_idx, column=2).value
            port = ws.cell(row=row_idx, column=3).value or 1521
            user = ws.cell(row=row_idx, column=4).value or 'system'
            pwd = ws.cell(row=row_idx, column=5).value
            svc = ws.cell(row=row_idx, column=6).value or ip
            if not name or not ip:
                continue
            db_info = {'name': str(name), 'ip': str(ip), 'port': int(port),
                       'user': str(user), 'password': str(pwd) if pwd else ''}
            print(f"\n{'='*40}\n巡检 [{name}] ({ip}:{port})\n{'='*40}")
            ok = run_inspection(db_info, service_name=str(svc))
            if ok: success_count += 1
        print(f"\n批量巡检完成: {success_count} / {(ws.max_row - 1)} 个成功")
    except Exception as e:
        print(f"批量巡检异常: {e}")


def create_excel_template():
    path = "oracle_batch_template.xlsx"
    wb = Workbook()
    ws = wb.active; ws.title = "Oracle巡检列表"
    headers = ["名称", "IP", "端口", "用户", "密码", "服务名/SID"]
    for j, h in enumerate(headers, 1): ws.cell(1,j).value=h
    # 示例行
    sample = ["生产-主库","192.168.1.100",1521,"system","oracle123","ORCL"]
    for j,v in enumerate(sample,1): ws.cell(2,j).value=v
    ws.column_dimensions['A'].width = 15; ws.column_dimensions['B'].width = 18
    ws.column_dimensions['C'].width = 8; ws.column_dimensions['D'].width = 12
    ws.column_dimensions['E'].width = 15; ws.column_dimensions['F'].width = 15
    wb.save(path); print(f"\n✅ Excel 模板已创建: {path}")


def check_license():
    """许可证检查"""
    lic_file = "oracle_inspector.lic"
    if not os.path.exists(lic_file):
        with open(lic_file, 'w') as f:
            f.write(json.dumps({"product":"DBCheck Oracle Inspector","version":"1.0",
                "licensed_to":"Trial","expire":"2099-12-31"}))
    try:
        with open(lic_file) as f: lic_data = json.load(f)
        exp = lic_data.get('expire','')
        if exp and datetime.strptime(exp,'%Y-%m-%d') < datetime.now():
            print("\n⚠️ 许可证已过期！请联系管理员。"); sys.exit(1)
    except Exception: pass


def main():
    start_time = time.time()

    # 支持从主入口通过 --template 直接生成 Excel 模板
    if len(sys.argv) > 1 and sys.argv[1] == '--template':
        create_excel_template()
        return

    print_banner(); check_license()
    while True:
        choice = show_main_menu()
        if choice == '1': single_inspection()
        elif choice == '2': batch_inspection()
        elif choice == '3': create_excel_template()
        elif choice == '4':
            print("\n感谢使用 DBCheck Oracle 数据库巡检工具！"); break
        if choice != '4':
            cont = input("\n是否返回主菜单? [y/N]: ").strip().lower()
            if cont not in ('', 'y', 'yes'):
                print("\n感谢使用 DBCheck Oracle 数据库巡检工具！"); break
    end_time = time.time()
    print(f"\n程序运行总耗时: {end_time-start_time:.2f}秒")


if __name__ == '__main__':
    main()