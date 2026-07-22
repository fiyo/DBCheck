#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
QA 真实端到端验证脚本（Issue4）：DB2 三修复 + SQL 编辑器底层链路

验证目标（真实 DB2：localhost:50000/testdb/db2inst1）：
  A. 数据源测试连接（对应 web_ui 数据源测试）        -> Issue? 连接可用性
  B. 完整巡检 + 报告生成
       · Issue2：采集进度日志 stdout 含 [DB2]
       · Issue1：报告文件名来自 i18n 键（非回退字面量），且 generate_report 产出正确文件名
       · Issue3：报告含「系统资源（CPU/内存/硬盘）」章节 + 磁盘表有数据行；context.system_info 含 cpu/memory/disk_list
  C. SQL 编辑器底层链路：DB-API cursor.execute 直连跑 DB2 专用查询，rows 非空

解释器：py -3.12 （DBCheck 根目录运行）
"""

import os
import sys
import io
import re
import traceback
import importlib.util
import contextlib

ROOT = os.path.dirname(os.path.abspath(__file__))
if ROOT not in sys.path:
    sys.path.insert(0, ROOT)

# ── 结果收集 ──
RESULTS = []

def check(name, cond, detail=""):
    status = "PASS" if cond else "FAIL"
    RESULTS.append((name, status, str(detail)))
    print(f"  [{status}] {name}" + (f"  -- {detail}" if detail else ""))

# ── 连接信息（用户提供的真实 DB2）──
DB_INFO = {
    'ip': 'localhost',
    'port': 50000,
    'database': 'testdb',
    'user': 'db2inst1',
    'password': 'password',
}

DB2_VERSION = 'unknown'
GENERATED_OFILE = None
COLLECTED_CONTEXT = {}
SQL_EDITOR_NOTE = ""

print("=" * 72)
print("A. 数据源测试连接（web_ui 数据源测试路径）")
print("=" * 72)
try:
    from plugin_loader import get_plugin_task_config
    cfg = get_plugin_task_config('db2')
    check("A: get_plugin_task_config('db2') 返回非 None", cfg is not None)
    if cfg is None:
        print("FATAL: 无法获取 db2 任务配置，终止。")
        sys.exit(1)
    check("A: cfg 含 connect_test 可调用", callable(cfg.get('connect_test')))
    check("A: cfg 含 connect_test_args 可调用", callable(cfg.get('connect_test_args')))

    ok, ver = cfg['connect_test'](*cfg['connect_test_args'](DB_INFO))
    DB2_VERSION = ver if ok else 'unknown'
    check("A: 连接测试 ok is True（localhost:50000/testdb/db2inst1）", ok is True, repr(ver))
    if ok:
        print(f"  >> DB2 版本: {ver}")
except Exception as e:
    check("A: 测试连接未抛异常", False, str(e))
    traceback.print_exc()

print("=" * 72)
print("B. 完整巡检 + 报告生成（Issue1 / Issue2 / Issue3）")
print("=" * 72)
data = None
try:
    # 动态导入插件模块（复刻 web_ui.run_inspection_task 的插件分发路径）
    plugin_path = cfg['plugin_path']
    main_file = cfg['main_file']
    spec = importlib.util.spec_from_file_location(
        "db2_main_verify", os.path.join(plugin_path, main_file))
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    check("B: 插件模块动态导入成功", True, main_file)
    check("B: 模块含 getData 函数", hasattr(mod, 'getData'))

    # 取参数 + 构造数据源
    pos, kw = cfg['getdata_args'](DB_INFO)
    data = mod.getData(*pos, **kw)
    check("B: getData 返回 CompatWrapper（非 None）", data is not None)

    # 捕获 stdout，验证 Issue2 进度日志
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        ctx = data.checkdb('builtin')
    stdout_text = buf.getvalue()
    COLLECTED_CONTEXT = ctx if isinstance(ctx, dict) else {}
    check("B: checkdb('builtin') 返回 context（dict）", isinstance(ctx, dict))
    # Issue2：采集进度日志应含 [DB2] 前缀（print_progress_bar 输出）
    check("B/Issue2: stdout 含 DB2 采集进度（'[DB2]'）",
          '[DB2]' in stdout_text,
          (stdout_text.strip().splitlines()[-3:] if stdout_text.strip() else ["(空)"]))
    print("  >> 采集进度片段（末 3 行）:")
    for line in stdout_text.strip().splitlines()[-3:]:
        print("     " + line[:130])
except Exception as e:
    check("B: 巡检采集未抛异常", False, str(e))
    traceback.print_exc()

print("-" * 72)
print("Issue1: 报告文件名模板（i18n 键解析，非回退字面量）")
print("-" * 72)
try:
    from i18n import t
    fname_tmpl = t(cfg['filename_key'])
    check("Issue1: i18n 键解析出真实模板（非回退字面量 webui.db2_report_filename）",
          fname_tmpl != 'webui.db2_report_filename',
          'template=' + fname_tmpl)
    rendered = fname_tmpl.format(ip='127.0.0.1', name='testdb', ts='20260722_120000')
    expected = 'DB2巡检报告_127.0.0.1_testdb_20260722_120000'
    check("Issue1: 文件名模板 .format(ip,name,ts) 正确",
          rendered == expected, rendered)

    reports_dir = os.path.join(ROOT, 'reports')
    os.makedirs(reports_dir, exist_ok=True)
    ofile = os.path.join(reports_dir, fname_tmpl + '.docx')
    GENERATED_OFILE = ofile
    if os.path.exists(ofile):
        os.remove(ofile)
    rc = data.generate_report(ofile, 'Jack')
    check("Issue1: generate_report 返回非 None", rc is not None)
    check("Issue1: 报告文件已生成", os.path.exists(ofile), ofile)
    check("Issue1: 文件名不是 webui.db2_report_filename.docx",
          not ofile.endswith('webui.db2_report_filename.docx'),
          os.path.basename(ofile))

    # 复刻 web_ui.py 真实命名路径（line 911-917, 923）：用 .format() 落真实值，
    # 证明生产环境报告文件名正确（非字面占位符）。
    import datetime as _dt
    ts_real = _dt.datetime.now().strftime('%Y%m%d_%H%M%S')
    label_name = DB_INFO.get('name', 'testdb') or DB_INFO['database']
    real_name = fname_tmpl.format(ip=DB_INFO['ip'], name=label_name, ts=ts_real)
    real_ofile = os.path.join(reports_dir, real_name + '.docx')
    if os.path.exists(real_ofile):
        os.remove(real_ofile)
    rc2 = data.generate_report(real_ofile, 'Jack')
    check("Issue1(生产路径): 真实格式化文件名 generate_report 成功",
          rc2 is not None and os.path.exists(real_ofile),
          os.path.basename(real_ofile))
    check("Issue1(生产路径): 文件名已落真实 ip/name/ts（无字面占位符）",
          ('{ip}' not in real_name) and ('{name}' not in real_name)
          and ('{ts}' not in real_name) and DB_INFO['ip'] in real_name,
          real_name)
except Exception as e:
    check("Issue1: 文件名/报告生成未抛异常", False, str(e))
    traceback.print_exc()

print("-" * 72)
print("Issue3: 系统资源章节（CPU / 内存 / 硬盘）")
print("-" * 72)
try:
    from docx import Document
    if os.path.exists(GENERATED_OFILE):
        doc = Document(GENERATED_OFILE)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        check("Issue3: 报告含『系统资源（CPU / 内存 / 硬盘）』标题",
              ('系统资源' in full_text and 'CPU' in full_text
               and '内存' in full_text and '磁盘' in full_text),
              "")
        # 查找磁盘表（5 列，表头含『挂载点』/『使用率(%)』）
        has_disk_table = False
        disk_data_rows = 0
        for tbl in doc.tables:
            headers = [c.text for c in tbl.rows[0].cells]
            if ('挂载点' in headers or '总容量' in headers
                    or '使用率(%)' in headers or 'Mount' in headers):
                has_disk_table = True
                disk_data_rows = max(disk_data_rows, len(tbl.rows) - 1)
        check("Issue3: 存在磁盘资源表", has_disk_table)
        check("Issue3: 磁盘表含数据行（>=1）", disk_data_rows >= 1,
              f"data_rows={disk_data_rows}")

        si = COLLECTED_CONTEXT.get('system_info') or {}
        check("Issue3: context.system_info 非空", bool(si), f"keys={list(si.keys())}")
        check("Issue3: context.system_info 含 cpu", bool(si.get('cpu')))
        check("Issue3: context.system_info 含 memory", bool(si.get('memory')))
        check("Issue3: context.system_info 含 disk_list", bool(si.get('disk_list')))
    else:
        check("Issue3: 报告文件存在可供检查", False, "ofile not found")
except Exception as e:
    check("Issue3: 章节检查未抛异常", False, str(e))
    traceback.print_exc()

print("=" * 72)
print("C. SQL 编辑器底层链路（DB-API cursor.execute 直连 DB2）")
print("=" * 72)
try:
    # 底层执行原语：JdbcCursorWrapper.execute + fetchall（SQL 编辑器对其它库型调同名 cursor.execute）
    inspector = mod.Db2JdbcInspector(
        DB_INFO['ip'], int(DB_INFO['port']), DB_INFO['user'], DB_INFO['password'],
        database=DB_INFO['database'])
    ok_c, ver_c = inspector.connect()
    check("C: 另建连接成功（供 SQL 执行验证）", ok_c, repr(ver_c))
    if ok_c:
        cur = inspector.conn.cursor()
        cur.execute("SELECT 1 FROM SYSIBM.SYSDUMMY1")
        rows = cur.fetchall()
        check("C: SELECT 1 FROM SYSIBM.SYSDUMMY1 返回非空 rows",
              isinstance(rows, list) and len(rows) >= 1, f"rows={rows}")
        check("C: 首行首列 == 1", bool(rows) and rows[0][0] == 1,
              str(rows[0] if rows else None))

        cur.execute("SELECT CURRENT TIMESTAMP FROM SYSIBM.SYSDUMMY1")
        rows2 = cur.fetchall()
        check("C: SELECT CURRENT TIMESTAMP FROM SYSIBM.SYSDUMMY1 返回非空 rows",
              isinstance(rows2, list) and len(rows2) >= 1, f"rows={rows2}")
        inspector.disconnect()

    # 静态核对 web_ui.py 的 SQL 编辑器路由是否含 db2 分支（已知增强项，不改判 Issue4）
    with open(os.path.join(ROOT, 'web_ui.py'), 'r', encoding='utf-8') as f:
        ui_src = f.read()
    m = re.search(r"def api_execute_sql\(.*?\):(.*?)\n(?:@app\.route|\ndef )",
                  ui_src, re.S)
    seg = m.group(1) if m else ""
    has_db2_branch = ("'db2'" in seg) or ('"db2"' in seg)
    if has_db2_branch:
        check("C(UI): web_ui api_execute_sql 已含 db2 分支", True)
        SQL_EDITOR_NOTE = "web_ui SQL 编辑器路由已直接支持 db2。"
    else:
        check("C(UI): web_ui api_execute_sql 含 db2 分支", False,
              "当前无 db2 分支（视为已知增强项，不影响底层链路判定）")
        SQL_EDITOR_NOTE = ("注意：web_ui.py 的 api_execute_sql / api_inspection_execute_sql "
                           "按 db_type 分发，目前没有 'db2' 分支（点击 DB2 数据源的 SQL 编辑器会返回"
                           "『不支持的数据库类型』）。但底层 DB-API 执行原语（cursor.execute + fetchall）"
                           "已验证对 DB2 可用；如需 UI 直连 DB2 SQL 编辑器，需为 web_ui 的 SQL 编辑器路由"
                           "补充 db2 分支（建议作为独立增强项，不阻塞本 Issue4 三处修复判定）。")
except Exception as e:
    check("C: SQL 执行未抛异常", False, str(e))
    traceback.print_exc()

print("=" * 72)
print("SUMMARY")
print("=" * 72)
passed = sum(1 for _, s, _ in RESULTS if s == "PASS")
failed = sum(1 for _, s, _ in RESULTS if s == "FAIL")
total = len(RESULTS)
print(f"  总断言: {total}  通过: {passed}  失败: {failed}")
if failed:
    print("  失败项:")
    for n, s, d in RESULTS:
        if s == "FAIL":
            print(f"    - {n}  ({d})")

# 路由判定
source_bug_names = [n for n, s, d in RESULTS if s == "FAIL"
                    and not n.startswith("C(UI)")]
if source_bug_names:
    ROUTING = "Engineer"
elif failed and all(n.startswith("C(UI)") for n, s, _ in RESULTS if s == "FAIL"):
    ROUTING = "NoOne（三修复通过；C(UI) 为已知增强项）"
else:
    ROUTING = "NoOne"
print(f"  路由判定: {ROUTING}")
print(f"  通过率: {passed}/{total} = {100.0*passed/total:.0f}%")
print(f"  DB2 版本: {DB2_VERSION}")
print(f"  报告绝对路径: {GENERATED_OFILE}")
