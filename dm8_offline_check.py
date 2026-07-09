#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#
# Copyright (c) 2025-2026 fiyo (Jack Ge) <sdfiyon@gmail.com>
#
# This file is part of DBCheck, an open-source database health inspection tool.
# DBCheck Professional — 专有商业软件，保留一切权利（Proprietary Software, All Rights Reserved）.
# See LICENSE for full license text.
#

"""
DM8 离线存储健康检查模块

本模块完全独立实现，不依赖任何第三方 DM8 存储格式逆向工程代码。
仅使用 DM8 官方公开文档中的通用存储概念和标准二进制文件分析技术。

检查范围：
1. 数据文件 (.DBF) 完整性 — 文件大小、页对齐、空页检测
2. 控制文件 (dm.ctl) 解析 — 提取文件路径、表空间映射
3. 系统表空间 (SYSTEM.DBF) 存在性检查
4. 文件与控制文件交叉校验
5. 页面级统计 — 全零页、异常页检测
6. 目录级诊断 — 缺失文件、重复文件、孤立文件

适用场景：DM8 数据库实例无法启动时，对底层数据文件进行离线健康评估。
"""

import os
import re
import json
import struct
import hashlib
import time
from pathlib import Path
from datetime import datetime
from collections import defaultdict

# ── DM8 存储常量（来自 DM8 官方公开文档）──────────────────────────
# DM8 支持的页大小选项（官方手册明确列出）
SUPPORTED_PAGE_SIZES = [4096, 8192, 16384, 32768]
DEFAULT_PAGE_SIZE = 8192

# DM8 数据文件扩展名
DBF_EXTENSIONS = ['.DBF', '.dbf']
# DM8 控制文件名模式
CTL_PATTERN = re.compile(r'dm\.ctl$', re.IGNORECASE)
# DM8 系统表空间文件
SYSTEM_DBF = 'SYSTEM.DBF'

# 健康等级
HEALTH_LEVELS = {
    'healthy':   {'label': '健康',   'color': 'green',  'score_range': (90, 100)},
    'warning':   {'label': '警告',   'color': 'yellow', 'score_range': (60, 89)},
    'critical':  {'label': '严重',   'color': 'red',    'score_range': (0, 59)},
    'unknown':   {'label': '未知',   'color': 'gray',   'score_range': (-1, -1)},
}

# 诊断严重级别
SEVERITY_INFO = 'info'
SEVERITY_WARN = 'warning'
SEVERITY_ERROR = 'error'
SEVERITY_FATAL = 'fatal'

# ── 数据块损坏类型（基于通用二进制分析，不依赖 DM8 专有页格式）──
# 以下检测仅使用"页内容明显异常"的通用信号，不读取任何 DM8 页头私有偏移、
# 不依赖 bic-dmdul 代码、不依赖达梦未公开的页格式知识，零侵权风险。
#   - ZERO_PAGE    : 整页字节全为 0x00
#   - CONSTANT_FILL: 整页为单一非全零字节（异常填充 / 磁盘坏道特征）
#   - TRUNCATED    : 文件末页字节数不足页大小（文件被截断）
BLOCK_ZERO = 'ZERO_PAGE'
BLOCK_CONSTANT = 'CONSTANT_FILL'
BLOCK_TRUNCATED = 'TRUNCATED'


class DM8DataFileInfo:
    """单个 DM8 数据文件的基本信息"""

    def __init__(self, file_path: Path):
        self.path = file_path
        self.name = file_path.name
        self.exists = file_path.exists()
        self.size = 0
        self.mtime = None
        self.readable = False
        self.page_size = 0  # 检测到的页大小，0=未检测
        self.total_pages = 0
        self.trailing_bytes = 0  # 文件大小 % 页大小的余数
        self.zero_pages = 0  # 全零页数量
        self.corrupt_blocks = []  # 坏块列表：[{file_name,file_path,page_no,file_offset,type,tablespace}]
        self.empty = False
        self.too_small = False  # 小于一个页大小
        self.header_hex = ''  # 首页头部前 64 字节的十六进制
        self.md5_prefix = ''  # 首页前 512 字节的 MD5（用于快速去重检测）

        if self.exists:
            try:
                stat = file_path.stat()
                self.size = stat.st_size
                self.mtime = datetime.fromtimestamp(stat.st_mtime)
                self.readable = os.access(str(file_path), os.R_OK)
                self.empty = (self.size == 0)
                self.too_small = (0 < self.size < 4096)
            except (OSError, PermissionError):
                pass

    def to_dict(self) -> dict:
        return {
            'name': self.name,
            'path': str(self.path),
            'exists': self.exists,
            'size': self.size,
            'size_human': _human_size(self.size),
            'readable': self.readable,
            'page_size': self.page_size,
            'total_pages': self.total_pages,
            'trailing_bytes': self.trailing_bytes,
            'zero_pages': self.zero_pages,
            'corrupt_blocks': self.corrupt_blocks,
            'empty': self.empty,
            'too_small': self.too_small,
            'mtime': self.mtime.isoformat() if self.mtime else None,
            'header_hex': self.header_hex,
        }


class DM8ControlFileInfo:
    """DM8 控制文件解析结果"""

    def __init__(self, file_path: Path):
        self.path = file_path
        self.name = file_path.name
        self.exists = file_path.exists()
        self.size = 0
        self.readable = False
        self.data_files = []  # 控制文件中引用的数据文件路径列表
        self.tablespace_map = {}  # {表空间名: [文件路径, ...]}
        self.raw_text = ''  # 提取出的可读文本

        if self.exists:
            try:
                stat = file_path.stat()
                self.size = stat.st_size
                self.readable = os.access(str(file_path), os.R_OK)
            except (OSError, PermissionError):
                pass

    def to_dict(self) -> dict:
        return {
            'name': self.name,
            'path': str(self.path),
            'exists': self.exists,
            'size': self.size,
            'size_human': _human_size(self.size),
            'readable': self.readable,
            'data_files': self.data_files,
            'tablespace_map': self.tablespace_map,
        }


class DM8Diagnostic:
    """单条诊断结果"""

    def __init__(self, severity: str, code: str, message: str,
                 file_name: str = '', detail: str = ''):
        self.severity = severity  # info / warning / error / fatal
        self.code = code  # 诊断码，如 FILE_MISSING, PAGE_MISALIGNED
        self.message = message
        self.file_name = file_name
        self.detail = detail

    @property
    def score_penalty(self) -> int:
        """该诊断项对健康分的扣分"""
        return {
            SEVERITY_INFO: 0,
            SEVERITY_WARN: 5,
            SEVERITY_ERROR: 15,
            SEVERITY_FATAL: 30,
        }.get(self.severity, 0)

    def to_dict(self) -> dict:
        return {
            'severity': self.severity,
            'code': self.code,
            'message': self.message,
            'file_name': self.file_name,
            'detail': self.detail,
        }


class DM8OfflineHealthChecker:
    """
    DM8 离线存储健康检查器

    用法：
        checker = DM8OfflineHealthChecker("/data/dm8/data")
        result = checker.run()
        print(result['health_score'])
        print(result['diagnostics'])
    """

    def __init__(self, db_dir: str, page_size: int = 0):
        """
        Args:
            db_dir: DM8 数据文件目录路径
            page_size: 页大小（0=自动检测，或指定 4096/8192/16384/32768）
        """
        self.db_dir = Path(db_dir)
        self.page_size = page_size if page_size in SUPPORTED_PAGE_SIZES else 0
        self.data_files: list[DM8DataFileInfo] = []
        self.control_files: list[DM8ControlFileInfo] = []
        self.diagnostics: list[DM8Diagnostic] = []
        self.corrupt_blocks = []  # 所有数据文件的坏块汇总
        self.scan_stats = {
            'start_time': None,
            'end_time': None,
            'duration': 0,
            'total_files_scanned': 0,
            'total_pages_scanned': 0,
            'total_data_size': 0,
        }

    def run(self) -> dict:
        """执行完整的离线健康检查，返回结果字典"""
        self.scan_stats['start_time'] = datetime.now()

        # Step 1: 验证目录
        if not self._validate_directory():
            return self._build_result()

        # Step 2: 发现文件
        self._discover_files()

        # Step 3: 检测页大小
        if not self.page_size:
            self._detect_page_size()

        # Step 4: 分析数据文件
        self._analyze_data_files()

        # Step 5: 解析控制文件
        self._parse_control_files()

        # Step 6: 交叉校验
        self._cross_validate()

        # Step 7: 检查系统表空间
        self._check_system_tablespace()

        # Step 8: 目录级诊断
        self._check_directory_level()

        self.scan_stats['end_time'] = datetime.now()
        self.scan_stats['duration'] = (
            self.scan_stats['end_time'] - self.scan_stats['start_time']
        ).total_seconds()

        return self._build_result()

    # ── Step 1: 验证目录 ──────────────────────────────────────────

    def _validate_directory(self) -> bool:
        """验证目标目录是否存在且可访问"""
        if not self.db_dir.exists():
            self._add_diag(SEVERITY_FATAL, 'DIR_NOT_FOUND',
                           f'目录不存在: {self.db_dir}')
            return False

        if not self.db_dir.is_dir():
            self._add_diag(SEVERITY_FATAL, 'DIR_NOT_DIRECTORY',
                           f'路径不是目录: {self.db_dir}')
            return False

        if not os.access(str(self.db_dir), os.R_OK):
            self._add_diag(SEVERITY_FATAL, 'DIR_NOT_READABLE',
                           f'目录不可读: {self.db_dir}')
            return False

        return True

    # ── Step 2: 发现文件 ──────────────────────────────────────────

    def _discover_files(self):
        """递归扫描目录，找到所有 .DBF 和 dm.ctl 文件"""
        for entry in self.db_dir.rglob('*'):
            if entry.is_file():
                ext = entry.suffix
                if ext in DBF_EXTENSIONS:
                    info = DM8DataFileInfo(entry)
                    self.data_files.append(info)
                    self.scan_stats['total_files_scanned'] += 1
                elif CTL_PATTERN.search(entry.name):
                    info = DM8ControlFileInfo(entry)
                    self.control_files.append(info)
                    self.scan_stats['total_files_scanned'] += 1

        if not self.data_files:
            self._add_diag(SEVERITY_WARN, 'NO_DBF_FILES',
                           '未找到任何 .DBF 数据文件')

        if not self.control_files:
            self._add_diag(SEVERITY_WARN, 'NO_CTL_FILE',
                           '未找到控制文件 dm.ctl')

    # ── Step 3: 检测页大小 ────────────────────────────────────────

    def _detect_page_size(self):
        """
        自动检测页大小。

        策略：统计每个标准页大小下，有多少数据文件能完美对齐（trailing_bytes=0）。
        选择对齐文件数最多的页大小；如果平局，选择较大的页大小。
        原理：8KB 页的文件能被 4KB 和 8KB 整除，但不能被 16KB 整除（除非恰好偶数页）。
        因此对齐文件最多的页大小中，最大的那个就是真实页大小。
        """
        non_empty = [f for f in self.data_files if f.exists and f.size > 0]
        if not non_empty:
            self.page_size = DEFAULT_PAGE_SIZE
            self._add_diag(SEVERITY_INFO, 'PAGE_SIZE_DEFAULT',
                           f'无法检测页大小，使用默认值 {DEFAULT_PAGE_SIZE} 字节')
            return

        # 统计每个页大小的对齐文件数
        alignment_counts = {}
        for ps in SUPPORTED_PAGE_SIZES:
            count = sum(1 for f in non_empty if f.size % ps == 0)
            alignment_counts[ps] = count

        # 找到对齐文件数最多的页大小
        max_count = max(alignment_counts.values())
        if max_count == 0:
            # 所有文件都无法被任何标准页大小整除——可能文件已损坏
            self.page_size = DEFAULT_PAGE_SIZE
            self._add_diag(SEVERITY_WARN, 'PAGE_SIZE_MISMATCH',
                           f'所有数据文件大小均无法被标准页大小整除，使用默认值 {DEFAULT_PAGE_SIZE}')
            return

        # 在对齐数最多的候选中，选择最大的页大小
        candidates = [ps for ps, cnt in alignment_counts.items() if cnt == max_count]
        self.page_size = max(candidates)

        self._add_diag(SEVERITY_INFO, 'PAGE_SIZE_DETECTED',
                       f'检测到页大小: {self.page_size} 字节 '
                       f'({max_count}/{len(non_empty)} 个文件对齐)')

    # ── Step 4: 分析数据文件 ──────────────────────────────────────

    def _analyze_data_files(self):
        """逐个分析数据文件的完整性"""
        ps = self.page_size or DEFAULT_PAGE_SIZE

        for info in self.data_files:
            self.scan_stats['total_data_size'] += info.size

            # 检查文件是否存在
            if not info.exists:
                self._add_diag(SEVERITY_ERROR, 'FILE_MISSING',
                               f'数据文件不存在: {info.name}',
                               file_name=info.name)
                continue

            # 检查文件是否为空
            if info.empty:
                self._add_diag(SEVERITY_ERROR, 'FILE_EMPTY',
                               f'数据文件为空 (0 字节): {info.name}',
                               file_name=info.name)
                continue

            # 检查文件是否过小
            if info.too_small:
                self._add_diag(SEVERITY_WARN, 'FILE_TOO_SMALL',
                               f'数据文件过小 ({info.size} 字节，小于最小页大小): {info.name}',
                               file_name=info.name)
                continue

            # 检查文件是否可读
            if not info.readable:
                self._add_diag(SEVERITY_ERROR, 'FILE_NOT_READABLE',
                               f'数据文件不可读: {info.name}',
                               file_name=info.name)
                continue

            # 计算页数和尾部字节
            info.page_size = ps
            info.total_pages = info.size // ps
            info.trailing_bytes = info.size % ps

            if info.trailing_bytes > 0:
                self._add_diag(SEVERITY_WARN, 'PAGE_MISALIGNED',
                               f'文件大小不是页大小的整数倍 (尾部 {info.trailing_bytes} 字节): {info.name}',
                               file_name=info.name,
                               detail=f'文件大小={info.size}, 页大小={ps}, 余数={info.trailing_bytes}')

            # 扫描页面统计
            self._scan_pages(info, ps)

        self.scan_stats['total_pages_scanned'] = sum(
            f.total_pages for f in self.data_files if f.total_pages > 0
        )

    def _scan_pages(self, info: DM8DataFileInfo, page_size: int):
        """
        扫描数据文件的所有页面，识别坏块（数据块损坏）。

        坏块识别基于"页内容明显异常"的通用信号，不读取任何 DM8 页头私有偏移、
        不依赖 bic-dmdul 代码、不依赖达梦未公开的页格式知识，零侵权：
          - ZERO_PAGE    : 整页字节全为 0x00
          - CONSTANT_FILL: 整页为单一非全零字节（异常填充 / 磁盘坏道特征）
          - TRUNCATED    : 文件末页字节数不足页大小（文件被截断）

        每个坏块记录物理页号（= file_offset // page_size，纯文件布局数学）
        与文件偏移，便于精确定位。
        """
        try:
            zero_count = 0
            constant_count = 0
            corrupt_blocks = []
            header_buf = None
            md5_buf = None
            pages_checked = 0

            # 多扫描一页以覆盖末页截断（文件大小非页大小整数倍）；
            # 限制单文件最大扫描页数，防止超大文件耗时过长
            max_pages = min(
                info.total_pages + (1 if info.trailing_bytes > 0 else 0),
                500000
            )

            with open(info.path, 'rb') as f:
                for page_idx in range(max_pages):
                    file_offset = page_idx * page_size
                    page_data = f.read(page_size)
                    if len(page_data) < page_size:
                        # 文件末页不足页大小 → 截断损坏
                        corrupt_blocks.append({
                            'file_name': info.name,
                            'file_path': str(info.path),
                            'page_no': page_idx,
                            'file_offset': file_offset,
                            'type': BLOCK_TRUNCATED,
                            'tablespace': '',
                        })
                        break

                    pages_checked += 1

                    # 检测全零页
                    if page_data == b'\x00' * page_size:
                        zero_count += 1
                        corrupt_blocks.append({
                            'file_name': info.name,
                            'file_path': str(info.path),
                            'page_no': page_idx,
                            'file_offset': file_offset,
                            'type': BLOCK_ZERO,
                            'tablespace': '',
                        })
                        continue

                    # 检测整页单一非全零字节（异常填充 / 磁盘坏道特征）
                    if page_data == page_data[:1] * page_size:
                        constant_count += 1
                        corrupt_blocks.append({
                            'file_name': info.name,
                            'file_path': str(info.path),
                            'page_no': page_idx,
                            'file_offset': file_offset,
                            'type': BLOCK_CONSTANT,
                            'tablespace': '',
                        })
                        continue

                    # 保存首页头部信息
                    if page_idx == 0:
                        header_buf = page_data[:64]
                        info.header_hex = header_buf.hex()
                        md5_buf = page_data[:512]
                        info.md5_prefix = hashlib.md5(md5_buf).hexdigest()

            info.zero_pages = zero_count
            info.total_pages = pages_checked  # 使用实际扫描的页数
            info.corrupt_blocks = corrupt_blocks

            # 全零页诊断（保留按文件详细诊断）
            if zero_count > 0:
                pct = (zero_count / pages_checked * 100) if pages_checked > 0 else 0
                if pct > 50:
                    self._add_diag(SEVERITY_ERROR, 'EXCESSIVE_ZERO_PAGES',
                                   f'全零页比例过高 ({zero_count}/{pages_checked} = {pct:.1f}%): {info.name}',
                                   file_name=info.name)
                elif pct > 10:
                    self._add_diag(SEVERITY_WARN, 'ZERO_PAGES_FOUND',
                                   f'发现全零页 ({zero_count}/{pages_checked} = {pct:.1f}%): {info.name}',
                                   file_name=info.name)
                else:
                    self._add_diag(SEVERITY_INFO, 'ZERO_PAGES_FOUND',
                                   f'少量全零页 ({zero_count}/{pages_checked} = {pct:.1f}%): {info.name}',
                                   file_name=info.name)

        except PermissionError:
            self._add_diag(SEVERITY_ERROR, 'FILE_READ_DENIED',
                           f'读取文件被拒绝: {info.name}',
                           file_name=info.name)
        except Exception as e:
            self._add_diag(SEVERITY_ERROR, 'FILE_READ_ERROR',
                           f'读取文件异常: {info.name} - {e}',
                           file_name=info.name)

    # ── Step 5: 解析控制文件 ──────────────────────────────────────

    def _parse_control_files(self):
        """
        解析 DM8 控制文件 (dm.ctl)。

        DM8 控制文件是二进制格式，但其中包含数据文件路径的文本字符串。
        本方法使用通用文本提取技术（非格式逆向），从二进制数据中提取
        .DBF 路径和可能的表空间名。
        """
        for ctl in self.control_files:
            if not ctl.exists or not ctl.readable:
                continue

            try:
                raw = ctl.path.read_bytes()

                # 从二进制数据中提取 ASCII/UTF-8 字符串
                # 控制文件中数据文件路径以 .DBF 结尾
                text_content = self._extract_strings(raw)

                # 提取 .DBF 文件路径
                dbf_pattern = re.compile(
                    rb'[\x20-\x7e]*\.DBF',
                    re.IGNORECASE
                )
                matches = dbf_pattern.findall(raw)

                seen = set()
                for m in matches:
                    path_str = m.decode('ascii', errors='ignore').strip()
                    if path_str and path_str not in seen:
                        seen.add(path_str)
                        ctl.data_files.append(path_str)

                # 尝试提取表空间名（路径中可能包含表空间信息）
                # DM8 控制文件中的路径格式通常为：/path/to/tablespace_name.DBF
                for path_str in ctl.data_files:
                    basename = os.path.basename(path_str)
                    # 去掉扩展名作为表空间名的候选
                    tbs_name = os.path.splitext(basename)[0]
                    if tbs_name and tbs_name.upper() != 'SYSTEM':
                        ctl.tablespace_map.setdefault(tbs_name, []).append(path_str)

                ctl.raw_text = text_content[:8192]  # 保留前 8KB 文本

            except PermissionError:
                self._add_diag(SEVERITY_ERROR, 'CTL_READ_DENIED',
                               f'读取控制文件被拒绝: {ctl.name}',
                               file_name=ctl.name)
            except Exception as e:
                self._add_diag(SEVERITY_ERROR, 'CTL_PARSE_ERROR',
                               f'解析控制文件异常: {ctl.name} - {e}',
                               file_name=ctl.name)

    def _extract_strings(self, raw: bytes, min_length: int = 4) -> str:
        """从二进制数据中提取可读字符串（通用方法，非格式特定）"""
        result = []
        current = []
        for byte in raw:
            if 0x20 <= byte <= 0x7e:
                current.append(chr(byte))
            else:
                if len(current) >= min_length:
                    result.append(''.join(current))
                current = []
        if len(current) >= min_length:
            result.append(''.join(current))
        return '\n'.join(result)

    # ── Step 6: 交叉校验 ─────────────────────────────────────────

    def _cross_validate(self):
        """交叉校验：控制文件引用的文件 vs 磁盘上的实际文件"""
        if not self.control_files or not self.data_files:
            return

        # 磁盘上实际存在的 .DBF 文件名集合（仅比较文件名，不比较完整路径）
        disk_file_names = {f.name.upper() for f in self.data_files if f.exists}

        # 控制文件中引用的所有文件名（提取 basename）
        ctl_referenced_names = set()
        for ctl in self.control_files:
            for ref_path in ctl.data_files:
                basename = os.path.basename(ref_path).upper()
                if basename.endswith('.DBF'):
                    ctl_referenced_names.add(basename)

        # 检查：控制文件引用了但磁盘上不存在的文件
        missing_from_disk = ctl_referenced_names - disk_file_names
        for missing in sorted(missing_from_disk):
            self._add_diag(SEVERITY_ERROR, 'CTL_REF_FILE_MISSING',
                           f'控制文件引用的文件在磁盘上不存在: {missing}')

        # 检查：磁盘上存在但控制文件未引用的文件
        disk_only = disk_file_names - ctl_referenced_names
        for extra in sorted(disk_only):
            if extra.upper() != SYSTEM_DBF.upper():
                self._add_diag(SEVERITY_INFO, 'ORPHAN_DBF_FILE',
                               f'数据文件未在控制文件中引用: {extra}')

    # ── Step 7: 检查系统表空间 ────────────────────────────────────

    def _check_system_tablespace(self):
        """检查 SYSTEM.DBF 系统表空间文件"""
        system_files = [f for f in self.data_files
                        if f.name.upper() == SYSTEM_DBF.upper()]

        if not system_files:
            self._add_diag(SEVERITY_ERROR, 'SYSTEM_DBF_MISSING',
                           f'系统表空间文件 {SYSTEM_DBF} 不存在')
            return

        if len(system_files) > 1:
            self._add_diag(SEVERITY_WARN, 'SYSTEM_DBF_DUPLICATE',
                           f'发现多个 {SYSTEM_DBF} 文件 ({len(system_files)} 个)，需要确认正确的系统表空间')

        for sf in system_files:
            if sf.empty:
                self._add_diag(SEVERITY_FATAL, 'SYSTEM_DBF_EMPTY',
                               f'系统表空间文件为空: {sf.name}')
            elif sf.size < self.page_size:
                self._add_diag(SEVERITY_FATAL, 'SYSTEM_DBF_TOO_SMALL',
                               f'系统表空间文件过小 ({sf.size} 字节): {sf.name}')

    # ── Step 8: 目录级诊断 ───────────────────────────────────────

    def _check_directory_level(self):
        """目录级整体诊断"""
        # 检查文件数量
        if len(self.data_files) == 0:
            self._add_diag(SEVERITY_FATAL, 'NO_DATA_FILES',
                           '目录中没有任何数据文件')
            return

        # 检查文件大小总和
        total_size = sum(f.size for f in self.data_files)
        if total_size == 0:
            self._add_diag(SEVERITY_FATAL, 'ALL_FILES_EMPTY',
                           '所有数据文件大小均为 0')
            return

        # 检查是否有同名文件（不同子目录）
        name_count = defaultdict(int)
        for f in self.data_files:
            name_count[f.name.upper()] += 1
        for name, count in name_count.items():
            if count > 1:
                self._add_diag(SEVERITY_WARN, 'DUPLICATE_FILE_NAME',
                               f'发现同名数据文件 ({count} 个): {name}')

        # 检查是否有临时文件残留
        for f in self.data_files:
            if 'temp' in f.name.lower() or 'tmp' in f.name.lower():
                self._add_diag(SEVERITY_INFO, 'TEMP_FILE_FOUND',
                               f'发现可能的临时文件: {f.name}')

        # 检查控制文件数量
        if len(self.control_files) > 1:
            self._add_diag(SEVERITY_WARN, 'MULTIPLE_CTL_FILES',
                           f'发现多个控制文件 ({len(self.control_files)} 个)，可能导致混淆')

    # ── 工具方法 ──────────────────────────────────────────────────

    def _add_diag(self, severity: str, code: str, message: str,
                  file_name: str = '', detail: str = ''):
        self.diagnostics.append(
            DM8Diagnostic(severity, code, message, file_name, detail)
        )

    def _calculate_health_score(self) -> tuple:
        """计算健康分数和等级"""
        base_score = 100
        total_penalty = sum(d.score_penalty for d in self.diagnostics)
        score = max(0, base_score - total_penalty)

        for level, info in HEALTH_LEVELS.items():
            lo, hi = info['score_range']
            if lo <= score <= hi:
                return score, level

        return score, 'unknown'

    def _build_file_to_ts_map(self) -> dict:
        """构建 文件名(大写) → 表空间 的反向映射（来自控制文件解析结果）"""
        fmap = {}
        for ctl in self.control_files:
            tbs_map = getattr(ctl, 'tablespace_map', {}) or {}
            for tbs_name, paths in tbs_map.items():
                for p in paths:
                    fmap[os.path.basename(p).upper()] = tbs_name
        return fmap

    def _summarize_corrupt_blocks(self):
        """
        汇总所有数据文件的坏块，填充所属表空间，并生成汇总诊断。

        不读取任何 DM8 页头私有偏移，仅依赖文件布局数学与已解析的控制文件
        表空间映射，零侵权。
        """
        fmap = self._build_file_to_ts_map()
        summary = {
            'total': 0,
            BLOCK_ZERO: 0,
            BLOCK_CONSTANT: 0,
            BLOCK_TRUNCATED: 0,
        }
        all_blocks = []
        for info in self.data_files:
            for cb in getattr(info, 'corrupt_blocks', []):
                cb = dict(cb)
                cb['tablespace'] = fmap.get(cb['file_name'].upper(), '未知')
                summary[cb['type']] = summary.get(cb['type'], 0) + 1
                summary['total'] += 1
                all_blocks.append(cb)

        all_blocks.sort(key=lambda x: (x['file_name'], x['page_no']))
        self.corrupt_blocks = all_blocks

        if summary['total'] > 0:
            total_pages = self.scan_stats.get('total_pages_scanned', 0) or 0
            pct = (summary['total'] / total_pages * 100) if total_pages > 0 else 0
            msg = (f'发现 {summary["total"]} 个可疑坏块 '
                   f'(全零 {summary[BLOCK_ZERO]}, 异常填充 {summary[BLOCK_CONSTANT]}, '
                   f'截断 {summary[BLOCK_TRUNCATED]}，损坏率 {pct:.2f}%)')
            if pct > 5 or summary[BLOCK_TRUNCATED] > 0:
                sev = SEVERITY_ERROR
            elif pct > 1:
                sev = SEVERITY_WARN
            else:
                sev = SEVERITY_INFO
            self._add_diag(sev, 'BLOCK_CORRUPT_SUMMARY', msg)

    def _build_result(self) -> dict:
        """构建最终结果字典"""
        self._summarize_corrupt_blocks()
        score, level = self._calculate_health_score()

        # 按严重级别排序诊断
        severity_order = {SEVERITY_FATAL: 0, SEVERITY_ERROR: 1,
                          SEVERITY_WARN: 2, SEVERITY_INFO: 3}
        sorted_diags = sorted(self.diagnostics,
                              key=lambda d: severity_order.get(d.severity, 99))

        # 统计各级别诊断数
        diag_stats = defaultdict(int)
        for d in self.diagnostics:
            diag_stats[d.severity] += 1

        return {
            'health_score': score,
            'health_level': level,
            'health_label': HEALTH_LEVELS.get(level, {}).get('label', '未知'),
            'db_dir': str(self.db_dir),
            'page_size': self.page_size or DEFAULT_PAGE_SIZE,
            'data_files': [f.to_dict() for f in self.data_files],
            'control_files': [c.to_dict() for c in self.control_files],
            'diagnostics': [d.to_dict() for d in sorted_diags],
            'diag_stats': dict(diag_stats),
            'scan_stats': {
                'start_time': self.scan_stats['start_time'].isoformat() if self.scan_stats['start_time'] else None,
                'end_time': self.scan_stats['end_time'].isoformat() if self.scan_stats['end_time'] else None,
                'duration': round(self.scan_stats['duration'], 3),
                'total_files_scanned': self.scan_stats['total_files_scanned'],
                'total_pages_scanned': self.scan_stats['total_pages_scanned'],
                'total_data_size': self.scan_stats['total_data_size'],
                'total_data_size_human': _human_size(self.scan_stats['total_data_size']),
            },
            'timestamp': datetime.now().isoformat(),
            'mode': 'local',
            'corrupt_blocks': self.corrupt_blocks,
            'corrupt_summary': {
                'total': len(self.corrupt_blocks),
                'zero': sum(1 for b in self.corrupt_blocks if b['type'] == BLOCK_ZERO),
                'constant_fill': sum(1 for b in self.corrupt_blocks if b['type'] == BLOCK_CONSTANT),
                'truncated': sum(1 for b in self.corrupt_blocks if b['type'] == BLOCK_TRUNCATED),
            },
        }


def _human_size(size: int) -> str:
    """将字节大小转换为人类可读格式"""
    if size == 0:
        return '0 B'
    units = ['B', 'KB', 'MB', 'GB', 'TB']
    idx = 0
    while size >= 1024 and idx < len(units) - 1:
        size /= 1024.0
        idx += 1
    return f'{size:.1f} {units[idx]}'


def generate_offline_report_text(result: dict) -> str:
    """将检查结果生成为纯文本报告"""
    lines = []
    lines.append('=' * 70)
    lines.append('DM8 离线存储健康检查报告')
    lines.append('=' * 70)
    lines.append(f'检查目录:   {result["db_dir"]}')
    lines.append(f'页大小:     {result["page_size"]} 字节')
    lines.append(f'检查时间:   {result["timestamp"]}')
    lines.append(f'耗时:       {result["scan_stats"]["duration"]} 秒')
    lines.append('')

    score = result['health_score']
    level = result['health_label']
    lines.append(f'健康分数:   {score}/100  [{level}]')
    lines.append('')

    diag_stats = result.get('diag_stats', {})
    lines.append(f'诊断统计:   fatal={diag_stats.get("fatal", 0)}, '
                 f'error={diag_stats.get("error", 0)}, '
                 f'warning={diag_stats.get("warning", 0)}, '
                 f'info={diag_stats.get("info", 0)}')
    lines.append('')

    lines.append('-' * 70)
    lines.append('数据文件清单:')
    lines.append('-' * 70)
    for f in result['data_files']:
        status = '✓' if f['exists'] and f['readable'] and not f['empty'] else '✗'
        lines.append(f'  {status} {f["name"]}')
        lines.append(f'      大小: {f["size_human"]}, 页数: {f["total_pages"]}, '
                      f'零页: {f["zero_pages"]}, 尾部字节: {f["trailing_bytes"]}')
    lines.append('')

    lines.append('-' * 70)
    lines.append('控制文件:')
    lines.append('-' * 70)
    for c in result['control_files']:
        lines.append(f'  {c["name"]} ({c["size_human"]})')
        for ref in c['data_files'][:20]:
            lines.append(f'      -> {ref}')
        if len(c['data_files']) > 20:
            lines.append(f'      ... 共 {len(c["data_files"])} 个引用')
    lines.append('')

    lines.append('-' * 70)
    lines.append('诊断详情:')
    lines.append('-' * 70)
    for d in result['diagnostics']:
        sev_label = {'fatal': '严重', 'error': '错误',
                     'warning': '警告', 'info': '信息'}.get(d['severity'], d['severity'])
        lines.append(f'  [{sev_label}] {d["code"]}: {d["message"]}')
        if d['detail']:
            lines.append(f'      详情: {d["detail"]}')
    lines.append('')
    lines.append('=' * 70)

    return '\n'.join(lines)


def generate_offline_report_word(result: dict, output_path: str = '') -> str:
    """
    将检查结果生成为 Word (.docx) 报告。

    Args:
        result: DM8OfflineHealthChecker.run() 返回的结果字典
        output_path: 输出文件路径，为空时自动生成临时文件

    Returns:
        生成的 .docx 文件路径
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    import tempfile

    doc = Document()

    # ── 辅助函数 ──────────────────────────────────────────────
    def _set_cell_bg(cell, hex_color):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def _set_font(run, name='微软雅黑', size=Pt(10.5), bold=False,
                  color=None, _qn=qn, _Pt=Pt, _RGBColor=RGBColor):
        run.font.name = name
        run.font.size = size
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        r = run._element
        r.rPr.rFonts.set(_qn('w:eastAsia'), name)

    def _add_heading_styled(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.font.size = Pt(14) if level == 1 else Pt(12)
        return h

    def _style_header_row(table, headers, bg_color='336699'):
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = h
            _set_cell_bg(cell, bg_color)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    _set_font(run, size=Pt(9), bold=True,
                              color=RGBColor(255, 255, 255))

    def _style_data_cell(cell, text, size=Pt(9)):
        cell.text = str(text) if text is not None else ''
        for p in cell.paragraphs:
            for run in p.runs:
                _set_font(run, size=size)

    # ── 封面 ──────────────────────────────────────────────────
    # Logo
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                             'dbcheck_logo.png')
    if os.path.exists(logo_path):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        logo_run.add_picture(logo_path, width=Cm(3.5))

    # 主标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('DM8 离线存储健康检查报告')
    _set_font(title_run, size=Pt(28), bold=True, color=RGBColor(15, 75, 135))

    # 副标题
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run('DM8 Offline Storage Health Check Report')
    _set_font(sub_run, name='Times New Roman', size=Pt(14),
              color=RGBColor(100, 100, 100))
    sub_run.font.italic = True

    doc.add_paragraph()

    # ── 封面信息表格 ──────────────────────────────────────────
    mode = result.get('mode', 'local')
    is_remote = (mode == 'remote')

    info_rows = [
        ('检查模式', 'SSH 远程检查' if is_remote else '本机检查'),
        ('DM8 数据文件目录', result.get('db_dir', '--')),
    ]

    if is_remote:
        info_rows.extend([
            ('SSH 主机地址', result.get('ssh_host', '--')),
            ('SSH 端口', str(result.get('ssh_port', 22))),
            ('SSH 用户名', result.get('ssh_user', '--')),
        ])

    info_rows.extend([
        ('页大小', f'{result.get("page_size", 0)} 字节'),
        ('检查时间', result.get('timestamp', '--')),
        ('检查耗时', f'{result.get("scan_stats", {}).get("duration", 0)} 秒'),
        ('健康分数', f'{result.get("health_score", 0)}/100  [{result.get("health_label", "--")}]'),
    ])

    info_table = doc.add_table(rows=len(info_rows), cols=2, style='Table Grid')
    for i, (label, value) in enumerate(info_rows):
        row = info_table.rows[i]
        _style_data_cell(row.cells[0], label)
        _set_cell_bg(row.cells[0], '336699')
        for p in row.cells[0].paragraphs:
            for run in p.runs:
                _set_font(run, size=Pt(10.5), bold=True,
                          color=RGBColor(255, 255, 255))
        _style_data_cell(row.cells[1], value)

    doc.add_page_break()

    # ── 第一章：检查摘要 ──────────────────────────────────────
    _add_heading_styled('第一章 检查摘要', level=1)

    score = result.get('health_score', 0)
    level = result.get('health_level', 'unknown')
    label = result.get('health_label', '未知')

    # 健康等级颜色映射
    score_colors = {
        'critical': RGBColor(220, 38, 38),
        'warning': RGBColor(234, 179, 8),
        'caution': RGBColor(249, 115, 22),
        'healthy': RGBColor(34, 197, 94),
        'unknown': RGBColor(100, 100, 100),
    }
    score_color = score_colors.get(level, RGBColor(100, 100, 100))

    p = doc.add_paragraph()
    run = p.add_run(f'健康分数: {score}/100')
    _set_font(run, size=Pt(20), bold=True, color=score_color)

    p = doc.add_paragraph()
    run = p.add_run(f'健康等级: {label}')
    _set_font(run, size=Pt(14), bold=True, color=score_color)

    # 诊断统计
    diag_stats = result.get('diag_stats', {})
    p = doc.add_paragraph()
    run = p.add_run(
        f'诊断统计 — 严重: {diag_stats.get("fatal", 0)}, '
        f'错误: {diag_stats.get("error", 0)}, '
        f'警告: {diag_stats.get("warning", 0)}, '
        f'信息: {diag_stats.get("info", 0)}'
    )
    _set_font(run, size=Pt(11))

    # 扫描统计
    scan_stats = result.get('scan_stats', {})
    p = doc.add_paragraph()
    run = p.add_run(
        f'扫描统计 — 文件数: {scan_stats.get("total_files_scanned", 0)}, '
        f'总页数: {scan_stats.get("total_pages_scanned", 0)}, '
        f'数据总量: {scan_stats.get("total_data_size_human", "0 B")}, '
        f'耗时: {scan_stats.get("duration", 0)} 秒'
    )
    _set_font(run, size=Pt(11))

    doc.add_paragraph()

    # ── 第二章：诊断详情 ──────────────────────────────────────
    _add_heading_styled('第二章 诊断详情', level=1)

    diagnostics = result.get('diagnostics', [])
    if diagnostics:
        sev_labels = {
            'fatal': '严重', 'error': '错误',
            'warning': '警告', 'info': '信息'
        }
        sev_colors = {
            'fatal': 'DC2626', 'error': 'EF4444',
            'warning': 'EAB308', 'info': '3B82F6'
        }

        tbl = doc.add_table(rows=1 + len(diagnostics), cols=4,
                            style='Table Grid')
        _style_header_row(tbl, ['严重级别', '诊断码', '描述', '详情'])

        for idx, d in enumerate(diagnostics, 1):
            row = tbl.rows[idx]
            sev = d.get('severity', 'info')
            _style_data_cell(row.cells[0], sev_labels.get(sev, sev))
            _set_cell_bg(row.cells[0], sev_colors.get(sev, '808080'))
            for p in row.cells[0].paragraphs:
                for run in p.runs:
                    _set_font(run, size=Pt(9), bold=True,
                              color=RGBColor(255, 255, 255))
            _style_data_cell(row.cells[1], d.get('code', ''))
            _style_data_cell(row.cells[2], d.get('message', ''))
            _style_data_cell(row.cells[3], d.get('detail', ''))
    else:
        p = doc.add_paragraph()
        run = p.add_run('未发现任何诊断问题。')
        _set_font(run, size=Pt(11))

    doc.add_paragraph()

    # ── 第三章：数据文件清单 ──────────────────────────────────
    _add_heading_styled('第三章 数据文件清单', level=1)

    data_files = result.get('data_files', [])
    if data_files:
        tbl = doc.add_table(rows=1 + len(data_files), cols=7,
                            style='Table Grid')
        _style_header_row(tbl, [
            '文件名', '大小', '页数', '零页数',
            '尾部字节', '状态', '修改时间'
        ])

        for idx, f in enumerate(data_files, 1):
            row = tbl.rows[idx]
            _style_data_cell(row.cells[0], f.get('name', ''))
            _style_data_cell(row.cells[1], f.get('size_human', ''))
            _style_data_cell(row.cells[2], f.get('total_pages', 0))
            _style_data_cell(row.cells[3], f.get('zero_pages', 0))
            _style_data_cell(row.cells[4], f.get('trailing_bytes', 0))

            # 状态列
            exists = f.get('exists', False)
            readable = f.get('readable', False)
            empty = f.get('empty', False)
            if not exists:
                status = '缺失'
                bg = 'DC2626'
            elif empty:
                status = '空文件'
                bg = 'EF4444'
            elif not readable:
                status = '不可读'
                bg = 'EF4444'
            else:
                status = '正常'
                bg = '22C55E'
            _style_data_cell(row.cells[5], status)
            _set_cell_bg(row.cells[5], bg)
            for p in row.cells[5].paragraphs:
                for run in p.runs:
                    _set_font(run, size=Pt(9), bold=True,
                              color=RGBColor(255, 255, 255))

            mtime = f.get('mtime', '')
            if mtime:
                # 截取日期部分
                mtime_str = str(mtime)[:19].replace('T', ' ')
            else:
                mtime_str = '--'
            _style_data_cell(row.cells[6], mtime_str)
    else:
        p = doc.add_paragraph()
        run = p.add_run('未找到任何数据文件。')
        _set_font(run, size=Pt(11))

    doc.add_paragraph()

    # ── 第四章：控制文件 ──────────────────────────────────────
    _add_heading_styled('第四章 控制文件', level=1)

    control_files = result.get('control_files', [])
    if control_files:
        for ctl in control_files:
            _add_heading_styled(
                f'{ctl.get("name", "未知")} ({ctl.get("size_human", "--")})',
                level=2
            )
            refs = ctl.get('data_files', [])
            if refs:
                tbl = doc.add_table(rows=1 + len(refs), cols=1,
                                    style='Table Grid')
                _style_header_row(tbl, ['引用的数据文件路径'])
                for idx, ref in enumerate(refs, 1):
                    _style_data_cell(tbl.rows[idx].cells[0], ref)

            tbs_map = ctl.get('tablespace_map', {})
            if tbs_map:
                _add_heading_styled('表空间映射', level=2)
                tbs_rows = sum(len(v) for v in tbs_map.values())
                tbs_tbl = doc.add_table(rows=1 + tbs_rows, cols=2,
                                        style='Table Grid')
                _style_header_row(tbs_tbl, ['表空间名', '数据文件路径'])
                r_idx = 1
                for tbs_name, paths in tbs_map.items():
                    for path in paths:
                        _style_data_cell(tbs_tbl.rows[r_idx].cells[0], tbs_name)
                        _style_data_cell(tbs_tbl.rows[r_idx].cells[1], path)
                        r_idx += 1
            doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        run = p.add_run('未找到控制文件。')
        _set_font(run, size=Pt(11))

    # ── 第五章：数据块损坏分析 ──────────────────────────────
    _add_heading_styled('第五章 数据块损坏分析', level=1)

    corrupt_blocks = result.get('corrupt_blocks', [])
    csummary = result.get('corrupt_summary', {})
    total_scanned = result.get('scan_stats', {}).get('total_pages_scanned', 0)

    if csummary:
        rate = (csummary.get('total', 0) / total_scanned * 100) if total_scanned > 0 else 0
        p = doc.add_paragraph()
        run = p.add_run(
            f'扫描总页数: {total_scanned}  |  可疑坏块: {csummary.get("total", 0)}  |  '
            f'损坏率: {rate:.2f}%'
        )
        _set_font(run, size=Pt(11))

        p = doc.add_paragraph()
        run = p.add_run(
            f'坏块类型分布 — 全零页: {csummary.get("zero", 0)}, '
            f'异常填充页: {csummary.get("constant_fill", 0)}, '
            f'截断页: {csummary.get("truncated", 0)}'
        )
        _set_font(run, size=Pt(11))

    if corrupt_blocks:
        # 坏块数量可能很大，最多展示前 500 个，避免文档过大
        show_blocks = corrupt_blocks[:500]
        tbl = doc.add_table(rows=1 + len(show_blocks), cols=5,
                            style='Table Grid')
        _style_header_row(tbl, [
            '数据文件', '物理页号', '文件偏移(字节)', '损坏类型', '所属表空间'
        ])

        type_labels = {
            BLOCK_ZERO: '全零页',
            BLOCK_CONSTANT: '异常填充页',
            BLOCK_TRUNCATED: '截断页',
        }
        type_colors = {
            BLOCK_ZERO: 'EF4444',
            BLOCK_CONSTANT: 'F59E0B',
            BLOCK_TRUNCATED: 'DC2626',
        }
        for idx, b in enumerate(show_blocks, 1):
            row = tbl.rows[idx]
            _style_data_cell(row.cells[0], b.get('file_name', ''))
            _style_data_cell(row.cells[1], b.get('page_no', ''))
            _style_data_cell(row.cells[2], b.get('file_offset', ''))
            btype = b.get('type', '')
            _style_data_cell(row.cells[3], type_labels.get(btype, btype))
            _set_cell_bg(row.cells[3], type_colors.get(btype, '808080'))
            for p in row.cells[3].paragraphs:
                for run in p.runs:
                    _set_font(run, size=Pt(9), bold=True,
                              color=RGBColor(255, 255, 255))
            _style_data_cell(row.cells[4], b.get('tablespace', '未知'))

        if len(corrupt_blocks) > len(show_blocks):
            p = doc.add_paragraph()
            run = p.add_run(
                f'（仅展示前 {len(show_blocks)} 个坏块，'
                f'共 {len(corrupt_blocks)} 个，完整清单见 JSON 结果）'
            )
            _set_font(run, size=Pt(9), color=RGBColor(120, 120, 120))
    else:
        p = doc.add_paragraph()
        run = p.add_run('未发现可疑坏块。')
        _set_font(run, size=Pt(11))

    # ── 报告尾部 ──────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f'报告生成时间: {result.get("timestamp", "--")}  |  '
        f'DBCheck DM8 离线存储健康检查  |  '
        f'检查模式: {"SSH 远程" if is_remote else "本机"}'
    )
    _set_font(run, size=Pt(8), color=RGBColor(150, 150, 150))

    # ── 保存 ──────────────────────────────────────────────────
    if not output_path:
        fd, output_path = tempfile.mkstemp(
            suffix='_dm8_offline_report.docx',
            prefix='dm8_'
        )
        os.close(fd)

    doc.save(output_path)
    return output_path

class DM8RemoteHealthChecker(DM8OfflineHealthChecker):
    """
    通过 SSH 连接远程服务器执行 DM8 离线存储健康检查。

    继承 DM8OfflineHealthChecker，重写文件操作方法为 SSH 远程命令。
    诊断逻辑（页大小检测、交叉校验、系统表空间检查等）完全复用父类。

    用法：
        checker = DM8RemoteHealthChecker(
            "/data/dm8/data",
            ssh_host="192.168.1.100",
            ssh_port=22,
            ssh_user="root",
            ssh_password="xxx",
        )
        result = checker.run()
    """

    def __init__(self, db_dir: str, ssh_host: str, ssh_port: int = 22,
                 ssh_user: str = 'root', ssh_password: str = '',
                 ssh_key_file: str = '', page_size: int = 0):
        """
        Args:
            db_dir: 远程服务器上的 DM8 数据文件目录路径
            ssh_host: SSH 主机地址
            ssh_port: SSH 端口
            ssh_user: SSH 用户名
            ssh_password: SSH 密码
            ssh_key_file: SSH 私钥文件路径
            page_size: 页大小（0=自动检测）
        """
        # 不调用父类的 Path(db_dir)，远程模式下 db_dir 是字符串
        self.db_dir_str = db_dir
        self.db_dir = None  # 远程模式不使用本地 Path
        self.page_size = page_size if page_size in SUPPORTED_PAGE_SIZES else 0
        self.data_files: list[DM8DataFileInfo] = []
        self.control_files: list[DM8ControlFileInfo] = []
        self.diagnostics: list[DM8Diagnostic] = []
        self.corrupt_blocks = []  # 所有数据文件的坏块汇总
        self.scan_stats = {
            'start_time': None,
            'end_time': None,
            'duration': 0,
            'total_files_scanned': 0,
            'total_pages_scanned': 0,
            'total_data_size': 0,
        }

        # SSH 连接信息
        self.ssh_host = ssh_host
        self.ssh_port = ssh_port
        self.ssh_user = ssh_user
        self.ssh_password = ssh_password
        self.ssh_key_file = ssh_key_file
        self._ssh_client = None
        self._remote_is_windows = False  # 远程系统类型

    # ── SSH 连接管理 ──────────────────────────────────────────────

    def _get_ssh_client(self):
        """建立并缓存 SSH 连接"""
        if self._ssh_client is not None:
            return self._ssh_client

        import paramiko
        client = paramiko.SSHClient()
        client.set_missing_host_key_policy(paramiko.AutoAddPolicy())

        connect_kwargs = {
            'hostname': self.ssh_host,
            'port': int(self.ssh_port),
            'username': self.ssh_user,
            'timeout': 15,
            'look_for_keys': False,
            'allow_agent': False,
        }

        if self.ssh_key_file and os.path.isfile(self.ssh_key_file):
            try:
                pkey = paramiko.RSAKey.from_private_key_file(self.ssh_key_file)
                connect_kwargs['pkey'] = pkey
            except Exception:
                try:
                    pkey = paramiko.Ed25519Key.from_private_key_file(self.ssh_key_file)
                    connect_kwargs['pkey'] = pkey
                except Exception:
                    pass

        if 'pkey' not in connect_kwargs and self.ssh_password:
            connect_kwargs['password'] = self.ssh_password

        client.connect(**connect_kwargs)
        self._ssh_client = client

        # 检测远程系统类型
        stdin, stdout, stderr = client.exec_command('uname -s 2>/dev/null || echo Windows')
        os_type = stdout.read().decode('utf-8', errors='ignore').strip()
        self._remote_is_windows = (os_type == 'Windows' or 'MINGW' in os_type)

        return client

    def _remote_exec(self, cmd: str, timeout: int = 30) -> tuple:
        """
        执行远程命令，返回 (stdout, stderr, exit_code)
        """
        client = self._get_ssh_client()
        stdin, stdout, stderr = client.exec_command(cmd, timeout=timeout)
        exit_code = stdout.channel.recv_exit_status()
        out = stdout.read().decode('utf-8', errors='ignore')
        err = stderr.read().decode('utf-8', errors='ignore')
        return out, err, exit_code

    def _remote_read_file_bytes(self, remote_path: str, max_bytes: int = 65536) -> bytes:
        """通过 SFTP 读取远程文件的前 N 字节"""
        client = self._get_ssh_client()
        sftp = client.open_sftp()
        try:
            with sftp.file(remote_path, 'rb') as f:
                return f.read(max_bytes)
        except Exception:
            return b''
        finally:
            sftp.close()

    def _remote_read_file_all(self, remote_path: str, max_bytes: int = 1048576) -> bytes:
        """通过 SFTP 读取远程文件全部内容（限制最大 1MB）"""
        return self._remote_read_file_bytes(remote_path, max_bytes)

    def close(self):
        """关闭 SSH 连接"""
        if self._ssh_client is not None:
            try:
                self._ssh_client.close()
            except Exception:
                pass
            self._ssh_client = None

    # ── 重写 Step 1: 验证目录 ────────────────────────────────────

    def _validate_directory(self) -> bool:
        """通过 SSH 验证远程目录是否存在且可访问"""
        try:
            if self._remote_is_windows:
                cmd = f'test -d "{self.db_dir_str}" && echo OK || echo NOTDIR'
            else:
                cmd = f'test -d "{self.db_dir_str}" && echo OK || echo NOTDIR'

            out, err, code = self._remote_exec(cmd)
            if 'OK' not in out:
                # 可能目录不存在
                out2, _, _ = self._remote_exec(f'test -e "{self.db_dir_str}" && echo EXISTS || echo NOTFOUND')
                if 'NOTFOUND' in out2:
                    self._add_diag(SEVERITY_FATAL, 'DIR_NOT_FOUND',
                                   f'远程目录不存在: {self.db_dir_str}')
                else:
                    self._add_diag(SEVERITY_FATAL, 'DIR_NOT_DIRECTORY',
                                   f'远程路径不是目录: {self.db_dir_str}')
                return False

            # 检查可读性
            out3, _, _ = self._remote_exec(f'test -r "{self.db_dir_str}" && echo READABLE || echo DENIED')
            if 'DENIED' in out3:
                self._add_diag(SEVERITY_FATAL, 'DIR_NOT_READABLE',
                               f'远程目录不可读: {self.db_dir_str}')
                return False

            return True

        except Exception as e:
            self._add_diag(SEVERITY_FATAL, 'SSH_CONNECTION_FAILED',
                           f'SSH 连接或命令执行失败: {e}')
            return False

    # ── 重写 Step 2: 发现文件 ──────────────────────────────────────

    def _discover_files(self):
        """通过 SSH find 命令递归扫描远程目录"""
        db_dir = self.db_dir_str

        # 使用 find + stat 一次性获取所有文件信息
        # 格式: path|size|mtime
        if self._remote_is_windows:
            # Windows 远程（Git Bash 环境）
            cmd = (
                f'find "{db_dir}" -type f '
                f'\\( -iname "*.DBF" -o -iname "dm.ctl" \\) '
                f'-exec stat --printf="%n|%s|%Y\\n" {{}} \\; 2>/dev/null'
            )
        else:
            cmd = (
                f'find "{db_dir}" -type f '
                f'\\( -iname "*.DBF" -o -iname "dm.ctl" \\) '
                f'-exec stat --printf="%n|%s|%Y\\n" {{}} \\; 2>/dev/null'
            )

        out, err, code = self._remote_exec(cmd, timeout=60)

        for line in out.strip().split('\n'):
            line = line.strip()
            if not line or '|' not in line:
                continue

            parts = line.split('|')
            if len(parts) < 3:
                continue

            file_path = parts[0]
            try:
                file_size = int(parts[1])
            except ValueError:
                file_size = 0
            try:
                file_mtime = float(parts[2])
            except ValueError:
                file_mtime = 0

            file_name = os.path.basename(file_path)
            ext = os.path.splitext(file_name)[1]

            if ext in DBF_EXTENSIONS:
                # 创建 DM8DataFileInfo 但不访问本地文件系统
                info = DM8DataFileInfo.__new__(DM8DataFileInfo)
                info.path = file_path  # 存储远程路径字符串
                info.name = file_name
                info.exists = True
                info.size = file_size
                info.mtime = datetime.fromtimestamp(file_mtime) if file_mtime else None
                info.readable = True  # find 能找到说明可读
                info.page_size = 0
                info.total_pages = 0
                info.trailing_bytes = 0
                info.zero_pages = 0
                info.empty = (file_size == 0)
                info.too_small = (0 < file_size < 4096)
                info.header_hex = ''
                info.md5_prefix = ''

                self.data_files.append(info)
                self.scan_stats['total_files_scanned'] += 1

            elif CTL_PATTERN.search(file_name):
                info = DM8ControlFileInfo.__new__(DM8ControlFileInfo)
                info.path = file_path
                info.name = file_name
                info.exists = True
                info.size = file_size
                info.readable = True
                info.data_files = []
                info.tablespace_map = {}
                info.raw_text = ''

                self.control_files.append(info)
                self.scan_stats['total_files_scanned'] += 1

        if not self.data_files:
            self._add_diag(SEVERITY_WARN, 'NO_DBF_FILES',
                           '远程目录中未找到任何 .DBF 数据文件')

        if not self.control_files:
            self._add_diag(SEVERITY_WARN, 'NO_CTL_FILE',
                           '远程目录中未找到控制文件 dm.ctl')

    # ── 重写 Step 4 中的页面扫描 ──────────────────────────────────

    def _scan_pages(self, info: DM8DataFileInfo, page_size: int):
        """
        远程页面扫描：通过 SSH 在远程执行 Python 脚本，识别坏块
        （全零页 / 整页单一字节异常填充 / 末页截断）。

        与本地版逻辑一致，仅检测"页内容明显异常"的通用信号，不读取 DM8 页头
        私有偏移、不依赖 bic-dmdul 代码，零侵权。

        若远程不支持 python3，降级为仅检查首页头（标记 REMOTE_SCAN_LIMITED）。
        """
        remote_path = info.path if isinstance(info.path, str) else str(info.path)

        # 远程 python3 脚本：扫描零页 + 异常填充 + 截断，输出 JSON
        scan_script = (
            "python3 -c \""
            "import sys,os,json,binascii;"
            "ps=" + str(page_size) + ";"
            "fp=sys.argv[1];"
            "total=os.path.getsize(fp);"
            "pages=total//ps;"
            "maxp=min(pages+(1 if total%ps else 0),500000);"
            "z=0;cf=0;cb=[];"
            "f=open(fp,'rb');"
            "hdr=f.read(ps)[:64];"
            "hx=binascii.hexlify(hdr).decode();"
            "zero=b'\\x00'*ps;"
            "for i in range(maxp):"
            " o=i*ps;"
            " d=f.read(ps);"
            " if len(d)<ps: cb.append({'page_no':i,'file_offset':o,'type':'TRUNCATED'});break;"
            " if d==zero: z+=1;cb.append({'page_no':i,'file_offset':o,'type':'ZERO_PAGE'});continue;"
            " if d==d[:1]*ps: cf+=1;cb.append({'page_no':i,'file_offset':o,'type':'CONSTANT_FILL'});"
            "f.close();"
            "print(json.dumps({'z':z,'pages':pages,'hx':hx,'cf':cf,'cb':cb}))"
            " \" \"" + remote_path + "\" 2>/dev/null"
        )

        out, err, code = self._remote_exec(scan_script, timeout=300)

        if code == 0 and out.strip():
            try:
                data = json.loads(out.strip().splitlines()[-1])
                info.zero_pages = int(data.get('z', 0))
                info.total_pages = int(data.get('pages', 0))
                info.header_hex = data.get('hx', '')
                info.page_size = page_size
                info.trailing_bytes = info.size % page_size
                # 还原坏块（tablespace 在汇总时填充）
                info.corrupt_blocks = [{
                    'file_name': info.name,
                    'file_path': str(info.path),
                    'page_no': b.get('page_no', 0),
                    'file_offset': b.get('file_offset', 0),
                    'type': b.get('type', ''),
                    'tablespace': '',
                } for b in data.get('cb', [])]
                self._diag_zero_pages(info)
                return
            except (ValueError, IndexError, json.JSONDecodeError):
                pass

        # 降级：通过 SFTP 读取首页头
        header_bytes = self._remote_read_file_bytes(remote_path, page_size)
        if header_bytes:
            info.header_hex = header_bytes[:64].hex()
            info.page_size = page_size
            info.total_pages = info.size // page_size
            info.trailing_bytes = info.size % page_size
            info.corrupt_blocks = []

            # 检查首页是否全零
            if len(header_bytes) >= page_size:
                if header_bytes[:page_size] == b'\x00' * page_size:
                    info.zero_pages = 1
                    info.corrupt_blocks = [{
                        'file_name': info.name,
                        'file_path': str(info.path),
                        'page_no': 0,
                        'file_offset': 0,
                        'type': BLOCK_ZERO,
                        'tablespace': '',
                    }]

            self._add_diag(SEVERITY_INFO, 'REMOTE_SCAN_LIMITED',
                           f'远程不支持 python3，仅检查了首页: {info.name}',
                           file_name=info.name)
        else:
            self._add_diag(SEVERITY_ERROR, 'FILE_READ_ERROR',
                           f'无法读取远程文件: {info.name}',
                           file_name=info.name)

    def _diag_zero_pages(self, info: DM8DataFileInfo):
        """全零页诊断（从父类 _scan_pages 中提取）"""
        zero_count = info.zero_pages
        pages_checked = info.total_pages
        if zero_count > 0:
            pct = (zero_count / pages_checked * 100) if pages_checked > 0 else 0
            if pct > 50:
                self._add_diag(SEVERITY_ERROR, 'EXCESSIVE_ZERO_PAGES',
                               f'全零页比例过高 ({zero_count}/{pages_checked} = {pct:.1f}%): {info.name}',
                               file_name=info.name)
            elif pct > 10:
                self._add_diag(SEVERITY_WARN, 'ZERO_PAGES_FOUND',
                               f'发现全零页 ({zero_count}/{pages_checked} = {pct:.1f}%): {info.name}',
                               file_name=info.name)
            else:
                self._add_diag(SEVERITY_INFO, 'ZERO_PAGES_FOUND',
                               f'少量全零页 ({zero_count}/{pages_checked} = {pct:.1f}%): {info.name}',
                               file_name=info.name)

    # ── 重写 Step 5: 解析控制文件 ──────────────────────────────────

    def _parse_control_files(self):
        """通过 SFTP 读取远程控制文件内容，然后本地解析"""
        for ctl in self.control_files:
            if not ctl.exists:
                continue

            remote_path = ctl.path if isinstance(ctl.path, str) else str(ctl.path)

            try:
                raw = self._remote_read_file_all(remote_path, max_bytes=1048576)

                if not raw:
                    self._add_diag(SEVERITY_ERROR, 'CTL_READ_ERROR',
                                   f'无法读取远程控制文件: {ctl.name}',
                                   file_name=ctl.name)
                    continue

                # 复用父类的解析逻辑
                text_content = self._extract_strings(raw)

                dbf_pattern = re.compile(
                    rb'[\x20-\x7e]*\.DBF',
                    re.IGNORECASE
                )
                matches = dbf_pattern.findall(raw)

                seen = set()
                for m in matches:
                    path_str = m.decode('ascii', errors='ignore').strip()
                    if path_str and path_str not in seen:
                        seen.add(path_str)
                        ctl.data_files.append(path_str)

                for path_str in ctl.data_files:
                    basename = os.path.basename(path_str)
                    tbs_name = os.path.splitext(basename)[0]
                    if tbs_name and tbs_name.upper() != 'SYSTEM':
                        ctl.tablespace_map.setdefault(tbs_name, []).append(path_str)

                ctl.raw_text = text_content[:8192]

            except Exception as e:
                self._add_diag(SEVERITY_ERROR, 'CTL_PARSE_ERROR',
                               f'解析远程控制文件异常: {ctl.name} - {e}',
                               file_name=ctl.name)

    # ── 重写 _build_result 添加远程信息 ────────────────────────────

    def _build_result(self) -> dict:
        """构建结果，添加远程模式标记和 SSH 连接信息"""
        result = super()._build_result()
        result['mode'] = 'remote'
        result['ssh_host'] = self.ssh_host
        result['ssh_port'] = self.ssh_port
        result['ssh_user'] = self.ssh_user
        result['db_dir'] = self.db_dir_str  # 覆盖为字符串路径
        return result


if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(description='DM8 离线存储健康检查工具')
    parser.add_argument('db_dir', help='DM8 数据文件目录路径')
    parser.add_argument('--page-size', type=int, default=0,
                        help='页大小 (4096/8192/16384/32768，0=自动检测)')
    parser.add_argument('--json', action='store_true', help='输出 JSON 格式')
    parser.add_argument('--remote', action='store_true', help='远程 SSH 模式')
    parser.add_argument('--ssh-host', default='', help='SSH 主机地址')
    parser.add_argument('--ssh-port', type=int, default=22, help='SSH 端口')
    parser.add_argument('--ssh-user', default='root', help='SSH 用户名')
    parser.add_argument('--ssh-password', default='', help='SSH 密码')
    parser.add_argument('--ssh-key', default='', help='SSH 私钥文件路径')
    args = parser.parse_args()

    if args.remote:
        if not args.ssh_host:
            print('Error: --ssh-host is required for remote mode')
            exit(1)
        checker = DM8RemoteHealthChecker(
            args.db_dir, args.ssh_host, args.ssh_port,
            args.ssh_user, args.ssh_password, args.ssh_key,
            args.page_size
        )
    else:
        checker = DM8OfflineHealthChecker(args.db_dir, args.page_size)

    result = checker.run()

    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print(generate_offline_report_text(result))

    if hasattr(checker, 'close'):
        checker.close()
